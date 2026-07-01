#!/usr/bin/env python3
"""Create Stage 23 pre-submission revision responding to reviewer-style critique."""

from __future__ import annotations

from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript"
INPUT_DOCX = MANUSCRIPT_DIR / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE22_BMC_COMPLIANT.docx"
OUTPUT_DOCX = MANUSCRIPT_DIR / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE23_PRE_SUBMISSION_REVISED.docx"
LOG_MD = MANUSCRIPT_DIR / "BMC_MIDM_STAGE23_PRE_SUBMISSION_REVISION_LOG.md"
MANUSCRIPT_TABLES = MANUSCRIPT_DIR / "tables"


SUPPLEMENTAL_TABLE_COPIES = {
    ROOT / "tables/table_uncertainty_triage.csv": MANUSCRIPT_TABLES / "supp_table_s4_uncertainty_triage.csv",
    ROOT / "tables/table_dca_summary.csv": MANUSCRIPT_TABLES / "supp_table_s5_dca_summary.csv",
    ROOT / "tables/table_unified_xai_cases.csv": MANUSCRIPT_TABLES / "supp_table_s6_xai_case_attribution_audit.csv",
}


def insert_paragraph_after(paragraph, text: str, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para.text = text
    if style is not None:
        new_para.style = style
    return new_para


def replace_exact(doc: Document, old: str, new: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == old.strip():
            paragraph.text = new
            return True
    return False


def replace_contains(doc: Document, old: str, new: str) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            count += 1
    return count


def find_paragraph(doc: Document, startswith: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Paragraph not found: {startswith}")


def apply_double_spacing(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 2
        paragraph.paragraph_format.space_after = Pt(0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_after = Pt(0)


def main() -> None:
    doc = Document(INPUT_DOCX)
    changes: list[str] = []
    MANUSCRIPT_TABLES.mkdir(parents=True, exist_ok=True)

    abstract_results_old = (
        "Results: In internal PTB-XL/PTB-XL+ testing, fair multimodal concatenation improved over "
        "the signal-embedding comparator in AUROC, average precision, and F1, with paired bootstrap "
        "confidence intervals excluding zero. Gated fusion showed only small numerical differences "
        "from fair concatenation, and paired bootstrap confidence intervals for AUROC, average precision, "
        "and F1 all contained zero. Signal-only external validation achieved macro AUROC 0.9071 on "
        "CPSC2018 and 0.8742 on Chapman-Shaoxing, with lower average precision and F1 in low-prevalence "
        "Chapman-Shaoxing labels. External calibration showed distribution-shift effects. The Stage 14L "
        "structured-feature audit found 138 allclose features, reduced structured-only internal collapse, "
        "no stable reduced-schema multimodal gain, and only two joinable external structured records per dataset."
    )
    abstract_results_new = (
        "Results: In internal PTB-XL/PTB-XL+ testing, fair multimodal concatenation achieved AUROC 0.9193, "
        "average precision 0.7953, and F1 0.7208, exceeding the signal-embedding comparator by +0.0098 AUROC, "
        "+0.0229 average precision, and +0.0205 F1 in paired bootstrap analysis. Gated fusion showed only small "
        "numerical differences from fair concatenation, with paired bootstrap confidence intervals containing zero. "
        "Signal-only external validation achieved macro AUROC 0.9071 on CPSC2018 and 0.8742 on Chapman-Shaoxing, "
        "but Chapman-Shaoxing AP and F1 were low in labels with low prevalence and transferred thresholds. "
        "External calibration showed distribution-shift effects. The structured-feature reproducibility audit found "
        "138 numerically concordant features, reduced-schema structured-only degradation, no stable reduced-schema "
        "multimodal gain, and only two joinable external structured records per dataset."
    )
    if replace_exact(doc, abstract_results_old, abstract_results_new):
        changes.append("Revised abstract Results to report internal fair-concat values and replace engineering terms.")

    replacements = {
        "The evaluation was staged. Stage 0/1 verified PTB-XL local data layout, parsed diagnostic labels, and generated frozen train, validation, and test files. Internal modeling stages trained and compared signal-only, signal-embedding, structured-only, fair concat, and gated fusion models. Later stages evaluated calibration, uncertainty triage, XAI, exploratory decision-curve analysis, signal-only external validation, and structured-feature reproducibility.": (
            "The evaluation followed a prespecified staged workflow. Initial data-ingestion checks verified the PTB-XL local data layout, parsed diagnostic labels, and generated frozen train, validation, and test files. Internal modeling then compared signal-only, signal-embedding, structured-only, fair concat, and gated fusion models. Subsequent analyses evaluated calibration, uncertainty triage, post-hoc explainability, exploratory decision-curve analysis, signal-only external validation, and structured-feature reproducibility."
        ),
        "Stage 14L structured-feature reproducibility audit": "Structured-feature reproducibility audit",
        "Stage 14L was treated as a core reproducibility and feasibility audit. Instead of continuing attempts to reconstruct the full PTB-XL+ 531-column schema externally, Stage 14L used a reduced structured-feature subset that had passed internal allclose checks in Stage 14H. This subset contained 138 features. The reduced schema was tested internally with signal-embedding, structured-only, and fair concat models, and external coverage was checked for CPSC2018 and Chapman-Shaoxing.": (
            "The structured-feature reproducibility analysis was treated as a core feasibility audit. Instead of continuing attempts to reconstruct the full PTB-XL+ 531-column schema externally, the audit used a reduced structured-feature subset that had passed internal numerical concordance checks against released PTB-XL+ values. This subset contained 138 features. The reduced schema was tested internally with signal-embedding, structured-only, and fair concat models, and external coverage was checked for CPSC2018 and Chapman-Shaoxing."
        ),
        "The go/no-go rule was conservative. External multimodal validation would only be considered if the reduced schema preserved internal multimodal gain and if external structured-feature coverage was adequate. If internal gain disappeared or external coverage was insufficient, external multimodal validation would remain NO-GO.": (
            "The decision rule was conservative. External multimodal validation would only be considered if the reduced schema preserved internal multimodal gain and if external structured-feature coverage was adequate. If internal gain disappeared or external coverage was insufficient, external multimodal validation would not be supported."
        ),
        "These results provide statistical support for the internal multimodal gain, while remaining limited to the PTB-XL/PTB-XL+ internal evaluation setting.": (
            "These results provide statistical support for an internal multimodal gain, while remaining limited to the PTB-XL/PTB-XL+ internal evaluation setting. The absolute gain was modest, particularly for AUROC, so the result should be interpreted as evidence of modality complementarity rather than as a large performance improvement."
        ),
        "Uncertainty triage, XAI, and exploratory decision-curve analysis were retained as decision-support diagnostics. These analyses broadened evaluation beyond discrimination, but they were not interpreted as clinical readiness evidence.": (
            "Uncertainty triage, post-hoc explainability, and exploratory decision-curve analysis were retained as decision-support diagnostics. In the gated-fusion model, retaining the lowest-uncertainty 80% of the test set retained 1,757 of 2,198 records and referred 441 records. In this retained subset, macro AUROC was 0.9303, AP was 0.8189, F1 was 0.7597, Brier score was 0.0687, and ECE was 0.0161, compared with overall macro AUROC 0.9196, AP 0.7978, F1 0.7314, Brier score 0.0844, and ECE 0.0193. Case-level explainability summaries identified retained correct examples, retained errors, and high-uncertainty referred errors; top structured features and signal leads were reported as audit outputs rather than causal explanations. Exploratory internal decision-curve analysis over thresholds 0.05 to 0.40 showed mean macro net benefit 0.1881 for fair concat, 0.1895 for gated fusion, 0.1844 for strong signal-only, and 0.1815 for structured-only models."
        ),
        "These results should be reported as transparent calibration behavior under distribution shift, not as evidence of clinical calibration readiness.": (
            "These results should be reported as transparent calibration behavior under distribution shift, not as evidence of clinical calibration readiness. The high external MCE values indicate that temperature scaling fitted on internal validation data did not transfer reliably to external label/data settings; therefore, external probabilities should not be interpreted as calibrated risks."
        ),
        "Stage 14L found that only 138 structured features were available in the reproducibility-validated allclose subset. In the reduced-schema internal test, the signal-embedding MLP achieved AUROC 0.9094, AP 0.7722, and F1 0.6981. The reduced structured-only MLP collapsed, with AUROC 0.5704, AP 0.3045, and F1 0.0000. Reduced fair concat achieved AUROC 0.9097, AP 0.7731, and F1 0.6938, which did not show stable gain over the signal-embedding MLP.": (
            "The structured-feature reproducibility audit found that only 138 structured features were available in the numerically concordant subset. In the reduced-schema internal test, the signal-embedding MLP achieved AUROC 0.9094, AP 0.7722, and F1 0.6981. The reduced structured-only MLP showed marked degradation, with AUROC 0.5704, AP 0.3045, and F1 0.0000. The zero F1 reflected thresholded predictions under the validation-derived thresholding rule after loss of most structured-feature information, rather than a successful reduced-schema classifier. Reduced fair concat achieved AUROC 0.9097, AP 0.7731, and F1 0.6938, which did not show stable gain over the signal-embedding MLP."
        ),
        "The Stage 14L decision was NO-GO for external multimodal validation. This should be presented as a structured-feature reproducibility and feasibility result, not as a failed external multimodal model.": (
            "The audit did not support external multimodal validation. This should be presented as a structured-feature reproducibility and feasibility result, not as a failed external multimodal model."
        ),
        "Stage 14L is a key BMC MIDM-compatible contribution because it shows why external multimodal claims are difficult in public multimodal ECG research. The audit found that a reproducibility-validated allclose subset contained only 138 features. This subset did not preserve internal multimodal gain, and external joinable structured coverage was only two records per external dataset. The practical implication is that released internal PTB-XL+ features can support internal multimodal experiments, but de novo reconstruction of the same structured schema on external WFDB datasets was not validated.": (
            "The structured-feature reproducibility audit is a key BMC MIDM-compatible contribution because it shows why external multimodal claims are difficult in public multimodal ECG research. The audit found that a reproducibility-validated numerically concordant subset contained only 138 features. This subset did not preserve internal multimodal gain, and external joinable structured coverage was only two records per external dataset. The practical implication is that released internal PTB-XL+ features can support internal multimodal experiments, but de novo reconstruction of the same structured schema on external WFDB datasets was not validated."
        ),
        "The study's main strength is transparent boundary-setting. It avoids claiming more than the data support and directly reports negative or near-null findings, including the gated-versus-concat null result and the Stage 14L external multimodal NO-GO decision. This approach is well suited to medical informatics, where model-performance claims need to be connected to reproducibility, calibration, external transportability, and auditability.": (
            "The study's main strength is transparent boundary-setting. It avoids claiming more than the data support and directly reports negative or near-null findings, including the gated-versus-concat null result and the unsupported external multimodal validation decision. This approach is well suited to medical informatics, where model-performance claims need to be connected to reproducibility, calibration, external transportability, and auditability."
        ),
        "Several limitations should be emphasized. First, all analyses used public retrospective datasets. Second, no prospective validation was performed. Third, external validation was signal-only; there was no validated external multimodal evaluation. Fourth, thresholds were not tuned on external test labels, which avoids overfitting but can reduce F1 under prevalence shift. Fifth, label mappings across PTB-XL, CPSC2018, and Chapman-Shaoxing were necessarily incomplete. Sixth, exact external reconstruction of PTB-XL+ structured features was not achieved. Seventh, decision-curve analysis was exploratory, internal, and threshold-dependent. Eighth, no clinical workflow, clinician interaction, or patient-outcome study was performed.": (
            "Several limitations should be emphasized. First, all analyses used public retrospective datasets. Second, no prospective validation was performed. Third, external validation was signal-only; there was no validated external multimodal evaluation. Fourth, thresholds were not tuned on external test labels, which avoids overfitting but can reduce F1 under prevalence shift. Fifth, label mappings across PTB-XL, CPSC2018, and Chapman-Shaoxing were necessarily incomplete, and very low-prevalence Chapman-Shaoxing labels should be treated as weak thresholded-performance evidence. Sixth, exact external reconstruction of PTB-XL+ structured features was not achieved. Seventh, decision-curve analysis was exploratory, internal, and threshold-dependent. Eighth, no clinical workflow, clinician interaction, or patient-outcome study was performed."
        ),
    }
    for old, new in replacements.items():
        if replace_exact(doc, old, new):
            changes.append(f"Replaced paragraph: {old[:60]}...")

    # Add statistical details without changing the model claims.
    stat_para = find_paragraph(doc, "Discrimination was summarized using macro AUROC")
    stat_para.text = (
        "Discrimination was summarized using macro AUROC, macro average precision, and macro F1. "
        "Macro metrics were computed as the unweighted mean across labels in the evaluated label scope. "
        "Thresholded metrics used thresholds selected on internal validation data. The frozen PTB-XL split "
        "included 17,418 training records, 2,183 validation records, and 2,198 test records. Fair concat was "
        "compared with the strongest unimodal comparators using paired record-level bootstrap resampling on the "
        "frozen internal test set. Gated fusion and fair concat were also compared using paired bootstrap "
        "resampling on the frozen internal test set [19]. Unless otherwise specified, bootstrap analyses used "
        "1,000 record-level resamples and percentile 95% confidence intervals. Bootstrap confidence intervals "
        "were reported without multiplicity adjustment and interpreted descriptively. External diagnostics were "
        "reported per dataset and per class. No external threshold optimization was performed."
    )
    changes.append("Added split sizes, macro-metric definition, bootstrap replicates, and CI method.")

    # Strengthen Chapman label/prevalence caveat in Results.
    chapman_para = find_paragraph(doc, "Chapman-Shaoxing signal-only external validation included")
    chapman_para.text = (
        "Chapman-Shaoxing signal-only external validation included 45,150 evaluated records for the MI/CD/HYP "
        "high-confidence label scope. Macro AUROC was 0.8742, macro AP was 0.1727, and macro F1 was 0.1650. "
        "The Chapman-Shaoxing results suggest preserved ranking performance but weak AP and F1. Per-class "
        "diagnostics support this interpretation: MI had 123 positive cases with prevalence 0.0027, AP 0.0796, "
        "and F1 0.0277; CD had prevalence 0.0677, AP 0.2624, and F1 0.3319; and HYP had prevalence 0.0166, "
        "AP 0.1760, and F1 0.1353. These values should be interpreted in relation to low prevalence, imperfect "
        "cross-dataset label comparability, and transfer of thresholds selected on PTB-XL validation data. The "
        "MI result in particular should be viewed as diagnostic context rather than strong evidence of comparable "
        "thresholded performance."
    )
    changes.append("Expanded Chapman-Shaoxing low-prevalence and label-comparability caveat.")

    # Strengthen discussion of practical effect size and calibration.
    principal_para = find_paragraph(doc, "The internal full-schema results showed")
    principal_para.text = (
        "The internal full-schema results showed that released PTB-XL+ structured features added value beyond "
        "ECG signal embeddings under frozen PTB-XL splits. Paired bootstrap analysis supported the fair concat "
        "gain over both the signal-embedding MLP and the strong signal-only comparator for AUROC, AP, and F1. "
        "The effect size was modest, however: the AUROC gain over the signal-embedding comparator was less than "
        "one percentage point, while AP and F1 gains were larger but still incremental. This distinction between "
        "statistical support and practical magnitude is important for decision-support interpretation. Gated fusion "
        "was numerically close to fair concat, but paired bootstrap intervals contained zero for AUROC, AP, and F1. "
        "This finding argues against making architectural complexity the central claim. For a decision-support "
        "evaluation paper, the more defensible contribution is the fair comparison framework and the statistically "
        "supported but modest internal modality complementarity."
    )
    changes.append("Added practical-magnitude interpretation for internal multimodal gain.")

    ext_disc_para = find_paragraph(doc, "The external evidence is intentionally limited")
    ext_disc_para.text = (
        "The external evidence is intentionally limited to signal-only validation. CPSC2018 supported limited "
        "signal-level external transportability for the NORM/CD high-confidence label subset. Chapman-Shaoxing "
        "showed preserved ranking performance but weak AP and F1, especially in low-prevalence labels. The "
        "per-class diagnostic table supports this explanation: Chapman-Shaoxing MI had only 123 positive cases "
        "among 45,150 evaluated records, with prevalence 0.0027, AP 0.0796, and F1 0.0277. This pattern is "
        "consistent with label mapping constraints, prevalence differences, and transfer of thresholds selected "
        "on internal validation data. It also means that the Chapman-Shaoxing MI result should be interpreted "
        "primarily as a stress test of signal ranking under label shift, not as robust thresholded-performance evidence. "
        "The absence of external threshold tuning strengthens the evaluation protocol but also limits thresholded performance."
    )
    changes.append("Added label-shift interpretation in Discussion.")

    # Supplementary material list additions.
    supp_fig = find_paragraph(doc, "Supplementary Figure S1:")
    cursor = insert_paragraph_after(supp_fig, "Supplementary Table S4: Uncertainty triage diagnostics.", supp_fig.style)
    cursor = insert_paragraph_after(cursor, "Supplementary Table S5: Exploratory decision-curve analysis.", supp_fig.style)
    insert_paragraph_after(cursor, "Supplementary Table S6: XAI case-selection and attribution audit.", supp_fig.style)
    changes.append("Added supplementary table pointers for uncertainty, DCA, and XAI.")

    # Clean remaining formal-language issues.
    for old, new in [
        ("Stage 14L", "structured-feature reproducibility audit"),
        ("allclose", "numerically concordant"),
        ("NO-GO", "not supported"),
    ]:
        n = replace_contains(doc, old, new)
        if n:
            changes.append(f"Replaced {n} remaining occurrence(s) of '{old}'.")

    for old, new in [
        (
            "structured-feature reproducibility audit structured-feature reproducibility audit",
            "structured-feature reproducibility audit",
        ),
        (
            "Supplementary Table S3: structured-feature reproducibility audit.",
            "Supplementary Table S3: Structured-feature reproducibility audit.",
        ),
        (
            "Supplementary Figure S1: structured-feature reproducibility audit.",
            "Supplementary Figure S1: Structured-feature reproducibility audit.",
        ),
        (
            "reduction from the released 531-column PTB-XL+ schema to the 138-feature numerically concordant subset",
            "reduction from the released 531-column PTB-XL+ schema to the 138-feature numerically concordant subset",
        ),
    ]:
        replace_contains(doc, old, new)

    apply_double_spacing(doc)
    doc.save(OUTPUT_DOCX)

    copied_tables: list[Path] = []
    for src, dst in SUPPLEMENTAL_TABLE_COPIES.items():
        copyfile(src, dst)
        copied_tables.append(dst)

    log = [
        "# BMC MIDM Stage 23 Pre-Submission Revision Log",
        "",
        "## Purpose",
        "",
        "This revision responds to reviewer-style pre-submission comments on the Stage 22 BMC-compliant manuscript.",
        "",
        "## Main Changes Implemented",
        "",
    ]
    for change in changes:
        log.append(f"- {change}")
    log.extend(
        [
            "",
            "## Evidence Boundary Preserved",
            "",
            "- Internal PTB-XL/PTB-XL+ full-schema multimodal evidence remains the only multimodal performance evidence.",
            "- External evidence remains signal-only for CPSC2018 and Chapman-Shaoxing.",
            "- The structured-feature reproducibility audit remains a feasibility/reproducibility finding, not an external multimodal result.",
            "- The revision does not claim gated-fusion superiority, exact external PTB-XL+ reconstruction, external multimodal validation, or clinical deployment readiness.",
            "",
            "## Items Deliberately Left For Later",
            "",
            "- A broader literature-expansion pass was not performed in this stage, because adding new recent ECG/multimodal references requires a separate citation-verification pass.",
            "- No new XAI figure was added; existing XAI outputs are now reported as supplementary audit tables because the previous figure was not readable at submission scale.",
            "",
            "## Generated Files",
            "",
            f"- Revised Word manuscript: `{OUTPUT_DOCX.relative_to(ROOT)}`",
            f"- Revision log: `{LOG_MD.relative_to(ROOT)}`",
        ]
    )
    for table_path in copied_tables:
        log.append(f"- Supplemental table copy: `{table_path.relative_to(ROOT)}`")
    LOG_MD.write_text("\n".join(log) + "\n", encoding="utf-8")
    print(OUTPUT_DOCX)
    print(LOG_MD)


if __name__ == "__main__":
    main()
