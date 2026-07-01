#!/usr/bin/env python3
"""Create Stage 27 manuscript revision after robustness audit."""

from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
IN_DOC = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE26_REVIEWER_RISK_REVISED.docx"
OUT_DOC = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE27_ROBUSTNESS_REVISED.docx"
LOG = ROOT / "manuscript" / "BMC_MIDM_STAGE27_ROBUSTNESS_REVISION_LOG.md"


def clear_and_write_labelled(paragraph, label: str, body: str) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    r = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    r.text = label
    r.bold = True
    body_run = paragraph.add_run(" " + body)
    body_run.bold = False


def set_text(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].bold = False
    else:
        paragraph.add_run(text)


def para_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.add_run(text)
    return inserted


def replace_para_contains(doc: Document, needle: str, new_text: str) -> None:
    for para in doc.paragraphs:
        if needle in para.text:
            set_text(para, new_text)
            return
    raise RuntimeError(f"Could not find paragraph containing: {needle}")


def copy_supplementary_tables() -> None:
    out_dir = ROOT / "manuscript" / "tables"
    out_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy2(ROOT / "tables" / "stage27_tolerance_sensitivity.csv", out_dir / "supp_table_s11_stage27_tolerance_sensitivity.csv")
    shutil.copy2(ROOT / "tables" / "stage27_feature_count_control_results.csv", out_dir / "supp_table_s12_stage27_feature_count_control_results.csv")
    shutil.copy2(ROOT / "tables" / "stage27_external_join_failure_audit.csv", out_dir / "supp_table_s13_stage27_external_join_failure_audit.csv")


def main() -> None:
    copy_supplementary_tables()
    shutil.copy2(IN_DOC, OUT_DOC)
    doc = Document(OUT_DOC)

    clear_and_write_labelled(
        doc.paragraphs[8],
        "Results:",
        (
            "With the released full PTB-XL+ schema, matched multimodal concatenation achieved AUROC 0.9193, average precision 0.7953, and F1 0.7208, exceeding the signal-embedding comparator by +0.0098 AUROC, +0.0229 average precision, and +0.0205 F1. "
            "Gated fusion showed no statistically clear advantage over concatenation. The structured-feature reproducibility audit found 138 numerically concordant features, but this concordant subset did not preserve internal multimodal gain and external structured-feature coverage was far below any usable threshold. "
            "Robustness controls showed that random-138 and internally importance-ranked 138-feature subsets recovered internal gains, indicating information loss in the concordance-selected subset rather than a 138-feature limit alone. "
            "Signal-only external validation achieved macro AUROC 0.9071 on CPSC2018 and 0.8742 on the Chapman-Shaoxing-Ningbo high-confidence label subset, with low AP/F1 in low-prevalence labels."
        ),
    )
    clear_and_write_labelled(
        doc.paragraphs[9],
        "Conclusions:",
        (
            "The main contribution is a reproducibility-aware evidence boundary: released PTB-XL+ features supported modest internal multimodal complementarity, but the current external structured-feature extraction-and-joining pipeline did not support external multimodal validation. "
            "Robustness checks indicated that the failed concordant subset reflected information loss from reproducibility-selected features, not feature count alone. "
            "Externally reproducible and sufficiently informative structured features are needed before external multimodal or clinical-use claims can be considered."
        ),
    )

    replace_para_contains(
        doc,
        "A tolerance-sensitivity analysis at looser numerical thresholds and a feature-count control analysis",
        (
            "The concordant-subset audit used a strict numerical agreement rule (absolute and relative tolerance 1e-6) to avoid treating implementation-dependent floating-point drift as interchangeable structured information. "
            "A Stage 27 robustness audit evaluated tolerance sensitivity and 138-feature count controls. The feature count remained 138 at 1e-4 and 1e-3 and increased only to 141 at 1e-2; substantially looser tolerances admitted more features but were not treated as exact reproduction. "
            "The same audit compared the concordance-selected 138 features with a random 138-feature subset and an internally attribution-ranked 138-feature subset under the same frozen internal split and validation-only model-selection rules. "
            "These controls were used to interpret the concordant-subset failure mechanism, not to establish external multimodal validation."
        ),
    )

    # Add Results paragraph after the Stage 14L reduced-schema paragraph.
    for para in doc.paragraphs:
        if "Reduced matched concat achieved AUROC 0.9097" in para.text:
            para_after(
                para,
                (
                    "A robustness audit clarified the mechanism of this reduced-schema failure. The strict concordance count was stable across tolerances up to 1e-3 and rose minimally to 141 features at 1e-2. "
                    "However, feature-count controls showed that 138 features were not inherently insufficient: a random 138-feature subset achieved matched-concat test AUROC 0.9195, AP 0.7960, and F1 0.7110, while an internally importance-ranked 138-feature subset achieved AUROC 0.9221, AP 0.8031, and F1 0.7310. "
                    "Both controls exceeded the signal-embedding comparator in screening paired bootstrap analyses using 300 record-level resamples. These findings indicate that the concordance-selected subset lost informative structured features; they do not support external multimodal validation because external candidate feature records remained limited to two per dataset."
                ),
            )
            break
    else:
        raise RuntimeError("Could not find Stage 14L reduced-schema result paragraph")

    replace_para_contains(
        doc,
        "The structured-feature reproducibility audit is a",
        (
            "The structured-feature reproducibility audit is an important BMC MIDM-compatible contribution because it separates internal reuse of released PTB-XL+ features from de novo reconstruction of comparable external structured features. "
            "The Stage 27 controls sharpened this interpretation. Random and internally importance-ranked 138-feature subsets recovered internal multimodal gains, whereas the concordance-selected 138-feature subset did not. "
            "Therefore, the audit should not be read as evidence that any 138-feature structured representation is inadequate. Rather, it shows that the features that were easiest to reproduce numerically were not the features carrying most of the internal multimodal signal, and that the current external extraction-and-joining pipeline did not yield enough candidate records for external multimodal testing. "
            "The practical implication remains conservative: released internal PTB-XL+ values can support internal multimodal experiments, but the present external WFDB reconstruction workflow is not sufficient for external multimodal validation."
        ),
    )

    # Update supplementary material list paragraph if present.
    for para in doc.paragraphs:
        if para.text.startswith("Supplementary Table S9"):
            current = para_after(para, "Supplementary Table S10: Reduced-schema internal threshold diagnostics.")
            current = para_after(current, "Supplementary Table S11: Stage 27 tolerance-sensitivity audit.")
            current = para_after(current, "Supplementary Table S12: Stage 27 feature-count control results.")
            para_after(current, "Supplementary Table S13: Stage 27 external join-failure audit.")
            break

    doc.save(OUT_DOC)

    tol = pd.read_csv(ROOT / "tables" / "stage27_tolerance_sensitivity.csv")
    res = pd.read_csv(ROOT / "tables" / "stage27_feature_count_control_results.csv")
    boot = pd.read_csv(ROOT / "tables" / "stage27_feature_count_control_bootstrap_ci.csv")
    default = res[(res["split"] == "test") & (res["threshold_mode"] == "default_0.5")]
    LOG.write_text(
        "# BMC MIDM Stage 27 Robustness Revision Log\n\n"
        "## Output\n"
        f"- Revised Word manuscript: `{OUT_DOC.relative_to(ROOT)}`\n"
        f"- Robustness audit summary: `docs/STAGE27_ROBUSTNESS_AUDIT.md`\n"
        "- Supplementary tables copied as S11-S13 under `manuscript/tables/`.\n\n"
        "## Key robustness findings\n"
        f"- Tolerance sensitivity: {int(tol.iloc[0]['n_features_passing'])} features at 1e-6; {int(tol[tol['atol'].eq(1e-3)].iloc[0]['n_features_passing'])} at 1e-3; {int(tol[tol['atol'].eq(1e-2)].iloc[0]['n_features_passing'])} at 1e-2.\n"
        "- Concordance-138 matched concat did not improve over the signal-embedding comparator.\n"
        "- Random-138 and internally importance-ranked 138 controls recovered internal multimodal gains.\n"
        "- External candidate structured records remained 2 per dataset; no external multimodal claim was added.\n\n"
        "## Manuscript interpretation change\n"
        "The manuscript now frames Stage 14L/27 as showing information loss in the numerically concordant subset, not as proof that a 138-feature representation is inherently inadequate.\n\n"
        "## Default-threshold internal control table\n\n"
        + default[["subset_name", "model", "auroc", "average_precision", "f1"]].to_markdown(index=False)
        + "\n\n## Bootstrap note\n"
        "Stage 27 paired bootstrap intervals use 300 record-level resamples as a lightweight screening audit; this is labelled in the summary and should not be confused with the primary manuscript bootstrap analyses.\n",
        encoding="utf-8",
    )
    print(OUT_DOC)
    print(LOG)


if __name__ == "__main__":
    main()
