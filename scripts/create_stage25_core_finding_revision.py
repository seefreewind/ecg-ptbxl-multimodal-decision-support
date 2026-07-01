#!/usr/bin/env python3
"""Create Stage 25 revision emphasizing the structured-feature reproducibility audit."""

from __future__ import annotations

import csv
from pathlib import Path
from shutil import copyfile

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript"
INPUT_DOCX = MANUSCRIPT_DIR / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE24_HARDENED.docx"
OUTPUT_DOCX = MANUSCRIPT_DIR / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE25_CORE_FINDING_REVISED.docx"
LOG_MD = MANUSCRIPT_DIR / "BMC_MIDM_STAGE25_CORE_FINDING_REVISION_LOG.md"
TABLES_DIR = MANUSCRIPT_DIR / "tables"
S8_INTERNAL_PER_CLASS = TABLES_DIR / "supp_table_s8_internal_per_superclass_metrics.csv"
S9_EXTERNAL_MAPPING = TABLES_DIR / "supp_table_s9_external_label_mapping_audit.csv"


def all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_exact(doc: Document, old: str, new: str) -> bool:
    for p in all_paragraphs(doc):
        if p.text.strip() == old.strip():
            p.text = new
            return True
    return False


def replace_contains(doc: Document, old: str, new: str) -> int:
    count = 0
    for p in all_paragraphs(doc):
        if old in p.text:
            p.text = p.text.replace(old, new)
            count += 1
    return count


def find_para(doc: Document, startswith: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(startswith)


def insert_after(paragraph, text: str, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para.text = text
    if style is not None:
        new_para.style = style
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def normalize_abstract_labels(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text
        for label in ("Background:", "Methods:", "Results:", "Conclusions:"):
            if text.startswith(label):
                rest = text[len(label):]
                p.clear()
                r = p.add_run(label)
                r.bold = True
                p.add_run(rest)
                break


def apply_spacing(doc: Document) -> None:
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 2
        p.paragraph_format.space_after = Pt(0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1
                    p.paragraph_format.space_after = Pt(0)


def remove_trailing_empty_paragraphs(doc: Document) -> int:
    removed = 0
    for p in reversed(doc.paragraphs):
        if p.text.strip():
            break
        remove_paragraph(p)
        removed += 1
    return removed


def make_supplementary_tables() -> None:
    TABLES_DIR.mkdir(parents=True, exist_ok=True)
    copyfile(ROOT / "figures/source_data/supplementary/supp_table_all_per_class_metrics.csv", S8_INTERNAL_PER_CLASS)
    mapping = pd.read_csv(ROOT / "tables/table_external_label_mapping_audit.csv")
    mapping["source_dataset"] = mapping["source_dataset"].replace({"Chapman-Shaoxing": "Chapman-Shaoxing-Ningbo"})
    mapping.to_csv(S9_EXTERNAL_MAPPING, index=False)


def main() -> None:
    make_supplementary_tables()
    doc = Document(INPUT_DOCX)
    changes: list[str] = []

    # Move the reproducibility audit into the abstract's foreground.
    replace_exact(
        doc,
        "Background: Multimodal electrocardiogram (ECG) models can combine waveform representations with structured ECG-derived measurements, but public-data studies need to distinguish internal multimodal performance from external transportability and feature-reproducibility constraints. We evaluated a conservative decision-support framework for ECG-based cardiac diagnostic assessment under explicit evidence boundaries.",
        "Background: Multimodal electrocardiogram (ECG) models can combine waveform representations with structured ECG-derived measurements, but public-data studies need to distinguish internal multimodal performance from external transportability and feature-reproducibility constraints. We evaluated whether released PTB-XL+ structured features support internal multimodal complementarity and whether comparable structured features can be reconstructed with enough fidelity and coverage for external multimodal validation.",
    )
    replace_exact(
        doc,
        "Methods: We used PTB-XL waveforms and released PTB-XL+ structured features for internal development and evaluation across five diagnostic superclasses: NORM, MI, STTC, CD, and HYP. Model groups included a strong signal-only model, a signal-embedding multilayer perceptron, a structured-feature multilayer perceptron, fair multimodal concatenation, and gated fusion. Model selection, threshold tuning, and temperature scaling used internal validation data only. We assessed discrimination, calibration, uncertainty triage, post-hoc explainability, and exploratory decision-curve analysis. External evaluation was restricted to signal-only validation on CPSC2018 and the Chapman-Shaoxing-Ningbo ECG database using restricted high-confidence label subsets. We separately audited whether structured PTB-XL+ compatible features could be reconstructed for external WFDB datasets.",
        "Methods: We used PTB-XL waveforms and released PTB-XL+ structured features for internal development and evaluation across five diagnostic superclasses: NORM, MI, STTC, CD, and HYP. Model groups included signal-only, signal-embedding, structured-feature, matched multimodal concatenation, and gated-fusion models. Model selection, threshold tuning, and temperature scaling used internal validation data only. External evaluation was restricted to signal-only validation on CPSC2018 and the Chapman-Shaoxing-Ningbo ECG database using restricted high-confidence label subsets. A structured-feature reproducibility audit tested whether a numerically concordant reduced feature subset could preserve internal multimodal gain and provide adequate external joinability.",
    )
    replace_exact(
        doc,
        "Results: In internal PTB-XL/PTB-XL+ testing, fair multimodal concatenation achieved AUROC 0.9193, average precision 0.7953, and F1 0.7208, exceeding the signal-embedding comparator by +0.0098 AUROC, +0.0229 average precision, and +0.0205 F1 in paired bootstrap analysis. Gated fusion showed only small numerical differences from fair concatenation, with paired bootstrap confidence intervals containing zero. Signal-only external validation achieved macro AUROC 0.9071 on CPSC2018 and 0.8742 on the Chapman-Shaoxing-Ningbo high-confidence label subset, but AP and F1 were low in labels with low prevalence and transferred thresholds. External calibration showed distribution-shift effects. The structured-feature reproducibility audit found 138 numerically concordant features, reduced-schema structured-only degradation, no stable reduced-schema multimodal gain, and only two joinable external structured records per dataset.",
        "Results: The structured-feature reproducibility audit found 138 numerically concordant features, reduced-schema structured-only degradation, no stable reduced-schema multimodal gain, and only two joinable external structured records per dataset; external multimodal validation was therefore not supported. In internal PTB-XL/PTB-XL+ testing with the released full structured schema, matched multimodal concatenation achieved AUROC 0.9193, average precision 0.7953, and F1 0.7208, exceeding the signal-embedding comparator by +0.0098 AUROC, +0.0229 average precision, and +0.0205 F1. Gated fusion showed no statistically clear advantage over concatenation. Signal-only external validation achieved macro AUROC 0.9071 on CPSC2018 and 0.8742 on the Chapman-Shaoxing-Ningbo high-confidence label subset, but AP and F1 were low in labels with low prevalence and transferred thresholds.",
    )
    replace_exact(
        doc,
        "Conclusions: The framework supports reproducibility-aware internal multimodal evaluation with signal-level external validation, but not external multimodal validation. Further work requires externally reproducible structured features and prospective clinical data before clinical-use claims can be considered.",
        "Conclusions: The main contribution is a reproducibility-aware evidence boundary: released PTB-XL+ features supported modest internal multimodal complementarity, but current external structured-feature extraction and joining did not support external multimodal validation. Further work requires externally reproducible structured features and prospective clinical data before clinical-use claims can be considered.",
    )
    changes.append("Rewrote abstract to foreground the structured-feature reproducibility audit.")

    # Introduction objective and principal finding.
    replace_exact(
        doc,
        "In this study, we evaluated a reproducibility-aware multimodal ECG decision-support framework using PTB-XL/PTB-XL+ internally and CPSC2018 and Chapman-Shaoxing externally. The objective was to test whether released PTB-XL+ structured features add internal value beyond ECG signal embeddings, whether simple fair concatenation is sufficient relative to gated fusion, how signal-only models behave on external datasets, and whether external structured features can be reconstructed with enough fidelity to support multimodal external validation. This was a conservative public-data evaluation, not a clinical deployment study.",
        "In this study, we evaluated a reproducibility-aware multimodal ECG decision-support framework using PTB-XL/PTB-XL+ internally and CPSC2018 plus the Chapman-Shaoxing-Ningbo ECG database externally. The primary objective was to define whether an internal PTB-XL/PTB-XL+ multimodal result can be responsibly extended to external multimodal validation when structured features must be reconstructed de novo. Secondary objectives were to quantify the internal value of released PTB-XL+ structured features beyond ECG signal embeddings, compare matched concatenation with gated fusion, and report signal-only external behavior under restricted high-confidence label subsets. This was a conservative public-data evaluation, not a clinical deployment study.",
    )
    replace_exact(
        doc,
        "This study evaluated a reproducibility-aware decision-support framework for multimodal ECG diagnostic classification. The main finding is not only that multimodal fusion improved internal PTB-XL/PTB-XL+ performance. The broader finding is that multimodal ECG models should be evaluated with explicit evidence boundaries: internal multimodal performance, signal-level external validation, calibration, uncertainty, explainability, decision-curve analysis, and structured-feature reproducibility should not be collapsed into one broad validation claim.",
        "This study evaluated a reproducibility-aware decision-support framework for multimodal ECG diagnostic classification. The principal finding is that internal multimodal evidence and external multimodal validation can diverge when structured features are released for the internal dataset but cannot be reconstructed and joined externally with adequate fidelity and coverage. The full-schema PTB-XL/PTB-XL+ experiments supported modest internal multimodal complementarity, but the structured-feature reproducibility audit did not support external multimodal validation. This evidence-boundary result is the central contribution of the study.",
    )
    changes.append("Moved the structured-feature audit to the introduction objective and Discussion principal finding.")

    # Controlled/matched language instead of fair as the main wording.
    replace_contains(doc, "fair multimodal concatenation", "matched multimodal concatenation")
    replace_contains(doc, "Fair concat", "Matched concat")
    replace_contains(doc, "fair concat", "matched concat")
    replace_contains(doc, "Fair MLP-concat", "Matched MLP-concat")
    replace_contains(doc, "fair MLP-concat", "matched MLP-concat")
    changes.append("Replaced most 'fair concat' wording with matched/controlled wording.")

    # Signal embedding source.
    model_para = find_para(doc, "The strong signal-only model used")
    model_para.text = model_para.text.replace(
        "The model provided both a direct waveform comparator and the 256-dimensional signal embedding for controlled multimodal comparison.",
        "The model provided both a direct waveform comparator and the 256-dimensional signal embedding for controlled multimodal comparison; embeddings were extracted from the penultimate representation of the trained signal model and then treated as frozen inputs for downstream MLP and matched-fusion experiments.",
    )
    insert_after(
        model_para,
        "The signal-embedding MLP and the strong signal-only model therefore answer related but distinct questions. The signal-embedding MLP controls for the frozen representation used by multimodal fusion, whereas the strong signal-only model evaluates the end-to-end waveform classifier. Both comparators are reported to avoid attributing a multimodal gain to classifier-head differences alone.",
        model_para.style,
    )
    changes.append("Clarified signal-embedding source and comparator rationale.")

    # Statistical details: bootstrap and macro DCA definition.
    stat = find_para(doc, "Discrimination was summarized using macro AUROC")
    stat.text = stat.text.replace(
        "Unless otherwise specified, bootstrap analyses used 1,000 record-level resamples and percentile 95% confidence intervals.",
        "For each bootstrap replicate, ECG records were resampled with replacement, all class-specific metrics were recomputed in the resampled set, macro metrics were then recomputed within that replicate, and model deltas were calculated on the paired resample. The primary internal multimodal-gain bootstrap used 2,000 record-level resamples; other diagnostic bootstrap summaries used 1,000 resamples and percentile 95% confidence intervals.",
    )
    insert_after(
        stat,
        "Exploratory decision-curve analysis was calculated per label and then summarized descriptively. The reported macro net benefit is the unweighted mean of label-specific net benefits at a given threshold; this aggregation is not a standard clinical utility endpoint and was used only to summarize internal threshold-dependent behavior. Label-specific DCA results are provided in supplementary source data.",
        stat.style,
    )
    changes.append("Clarified bootstrap resampling unit/aggregation and macro DCA definition.")

    # Per-class internal gain.
    internal_gain = find_para(doc, "These results provide statistical support")
    insert_after(
        internal_gain,
        "Per-superclass internal diagnostics showed that the largest matched-concat improvement over the signal-embedding comparator occurred for HYP, where AUROC increased from 0.8359 to 0.8791, AP from 0.4839 to 0.5918, and F1 from 0.4444 to 0.5393. Differences for NORM, MI, STTC, and CD were smaller and should be interpreted as descriptive per-class context rather than separate confirmatory tests.",
        internal_gain.style,
    )
    changes.append("Added per-superclass internal metric context and copied the full per-class table.")

    # Reduced structured-only F1: add validation-tuned value.
    reduced = find_para(doc, "The structured-feature reproducibility audit found")
    reduced.text = reduced.text.replace(
        "The zero F1 reflected thresholded predictions under the validation-derived thresholding rule after loss of most structured-feature information, rather than a successful reduced-schema classifier.",
        "The zero F1 reflected default-threshold predictions after loss of most structured-feature information; with validation-tuned thresholds, the same reduced structured-only model reached test F1 0.4002, confirming that the zero default-threshold F1 was a thresholding/probability-scale artifact rather than an empty evaluation pipeline.",
    )
    changes.append("Added validation-tuned reduced structured-only F1.")

    # External label mapping table and external wording.
    ext = find_para(doc, "CPSC2018 and the PhysioNet")
    insert_after(
        ext,
        "The external label mappings are provided as a supplementary audit table. Main external evaluation retained only high-confidence mappings: CPSC2018 contributed NORM and CD, while the Chapman-Shaoxing-Ningbo database contributed MI, CD, and HYP. Medium-confidence labels such as ST-T abnormalities were not used in the main external analysis. This design makes external results label-scope-specific and prevents direct comparison with the five-label internal macro metrics.",
        ext.style,
    )
    changes.append("Added external label-mapping methods text and supplementary mapping table.")

    # F1 and bootstrap footnote.
    table2 = find_para(doc, "Table 2 Statistical support")
    insert_after(
        table2,
        "Note. Deltas are calculated from unrounded prediction-level estimates and paired bootstrap resamples, so displayed deltas may differ slightly from subtraction of rounded table values.",
        table2.style,
    )
    f1_note = find_para(doc, "Uncertainty triage, post-hoc explainability")
    f1_note.text = f1_note.text.replace(
        "the primary Table 1 thresholded gated-fusion F1 remains 0.7255 under the main performance-evaluation pipeline.",
        "the primary Table 1 default-threshold gated-fusion F1 remains 0.7255 under the main performance-evaluation pipeline.",
    )
    changes.append("Added note about rounded point estimates versus bootstrap deltas and clarified F1 pipelines.")

    # Temperature scaling external no-refit.
    cal = find_para(doc, "External calibration was evaluated without external refitting")
    cal.text = cal.text + " This no-refit design was intentional: refitting temperature on external test labels would have violated the no-external-tuning evaluation principle."
    changes.append("Clarified no-refit external temperature scaling design.")

    # Supplementary list additions.
    s7 = find_para(doc, "Supplementary Table S7:")
    cursor = insert_after(s7, "Supplementary Table S8: Internal per-superclass performance diagnostics.", s7.style)
    insert_after(cursor, "Supplementary Table S9: External label-mapping audit.", s7.style)
    changes.append("Added S8/S9 supplementary table entries.")

    # TRIPOD+AI citation and remove unused PhysioNet/CinC 2020 reference.
    ai = find_para(doc, "AI-assisted programming")
    insert_after(ai, "A TRIPOD+AI-oriented reporting checklist is provided as Supplementary Table S7 [20].", ai.style)
    replace_contains(doc, "Residual learning was used as the general architectural background [17]", "Residual learning was used as the general architectural background [16]")
    replace_contains(doc, "medical AI cautions [9–10,18–19]", "medical AI cautions [9–10,17–18]")
    replace_contains(doc, "frozen internal test set [20]", "frozen internal test set [19]")

    # Remove old reference 16 and renumber following references.
    for p in list(doc.paragraphs):
        if p.text.startswith("16. Alday EAP"):
            remove_paragraph(p)
            break
    replace_exact(doc, "17. He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016:770–778.", "16. He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016:770–778.")
    replace_exact(doc, "18. Selvaraju RR, Cogswell M, Das A, Vedantam R, Parikh D, Batra D. Grad-CAM: Visual explanations from deep networks via gradient-based localization. IEEE International Conference on Computer Vision. 2017:618–626.", "17. Selvaraju RR, Cogswell M, Das A, Vedantam R, Parikh D, Batra D. Grad-CAM: Visual explanations from deep networks via gradient-based localization. IEEE International Conference on Computer Vision. 2017:618–626.")
    replace_exact(doc, "19. Ghassemi M, Oakden-Rayner L, Beam AL. The false hope of current approaches to explainable artificial intelligence in health care. The Lancet Digital Health. 2021;3:e745–e750. doi:10.1016/S2589-7500(21)00208-9.", "18. Ghassemi M, Oakden-Rayner L, Beam AL. The false hope of current approaches to explainable artificial intelligence in health care. The Lancet Digital Health. 2021;3:e745–e750. doi:10.1016/S2589-7500(21)00208-9.")
    replace_exact(doc, "20. Efron B, Tibshirani RJ. An Introduction to the Bootstrap. Chapman and Hall/CRC. 1993.", "19. Efron B, Tibshirani RJ. An Introduction to the Bootstrap. Chapman and Hall/CRC. 1993.")
    ref19 = find_para(doc, "19. Efron B")
    insert_after(ref19, "20. Collins GS, Moons KGM, Dhiman P, Riley RD, Beam AL, Van Calster B, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378. doi:10.1136/bmj-2023-078378.", ref19.style)
    changes.append("Removed unused PhysioNet/CinC 2020 reference and added TRIPOD+AI citation.")

    # Dataset naming inside embedded tables and manuscript-table CSV copies.
    replace_contains(doc, "Chapman-Shaoxing", "Chapman-Shaoxing-Ningbo")
    replace_contains(doc, "Chapman-Shaoxing-Ningbo-Ningbo", "Chapman-Shaoxing-Ningbo")
    replace_contains(doc, "10,646-record Chapman-Shaoxing-Ningbo publication", "10,646-record Chapman-Shaoxing publication")
    replace_contains(doc, "chapman", "Chapman-Shaoxing-Ningbo")
    for csv_path in [
        MANUSCRIPT_DIR / "tables/supp_table_external_per_class_diagnostics.csv",
        MANUSCRIPT_DIR / "tables/supp_table_external_per_class_diagnostics.md",
    ]:
        if csv_path.exists():
            data = csv_path.read_text(encoding="utf-8")
            data = data.replace("Chapman-Shaoxing", "Chapman-Shaoxing-Ningbo")
            data = data.replace("Chapman-Shaoxing-Ningbo-Ningbo", "Chapman-Shaoxing-Ningbo")
            data = data.replace("chapman", "Chapman-Shaoxing-Ningbo")
            csv_path.write_text(data, encoding="utf-8")

    normalize_abstract_labels(doc)
    removed = remove_trailing_empty_paragraphs(doc)
    if removed:
        changes.append(f"Removed {removed} trailing blank paragraph(s).")
    apply_spacing(doc)
    doc.save(OUTPUT_DOCX)

    log = [
        "# BMC MIDM Stage 25 Core Finding Revision Log",
        "",
        "## Purpose",
        "",
        "This revision foregrounds the structured-feature reproducibility audit as the central contribution and closes reviewer-facing gaps around per-class metrics, label mapping, bootstrap interpretation, and reporting guidance.",
        "",
        "## Changes",
        "",
    ]
    log.extend(f"- {c}" for c in changes)
    log.extend(
        [
            "",
            "## Generated Files",
            "",
            f"- `{OUTPUT_DOCX.relative_to(ROOT)}`",
            f"- `{LOG_MD.relative_to(ROOT)}`",
            f"- `{S8_INTERNAL_PER_CLASS.relative_to(ROOT)}`",
            f"- `{S9_EXTERNAL_MAPPING.relative_to(ROOT)}`",
        ]
    )
    LOG_MD.write_text("\n".join(log) + "\n", encoding="utf-8")
    print(OUTPUT_DOCX)
    print(LOG_MD)
    print(S8_INTERNAL_PER_CLASS)
    print(S9_EXTERNAL_MAPPING)


if __name__ == "__main__":
    main()
