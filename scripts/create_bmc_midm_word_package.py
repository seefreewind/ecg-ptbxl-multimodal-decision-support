#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_MD = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT.md"
COVER_MD = ROOT / "manuscript" / "BMC_MIDM_COVER_LETTER_DRAFT.md"
BIB = ROOT / "manuscript" / "references.bib"
OUT_DOCX = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT.docx"
COVER_DOCX = ROOT / "manuscript" / "BMC_MIDM_COVER_LETTER_DRAFT.docx"


FIGURES = [
    (
        "Figure 1. Study design and evidence boundaries.",
        ROOT / "manuscript" / "figures" / "figure1_study_design_evidence_boundary.png",
    ),
    (
        "Figure 2. Internal model performance and statistically supported multimodal gain.",
        ROOT / "manuscript" / "figures" / "figure2_internal_model_performance.png",
    ),
    (
        "Figure 3. Signal-only external validation.",
        ROOT / "manuscript" / "figures" / "figure3_external_signal_only_validation.png",
    ),
    (
        "Figure 4. Calibration and reliability under internal testing and external distribution shift.",
        ROOT / "manuscript" / "figures" / "figure4_calibration_reliability.png",
    ),
    (
        "Supplementary Figure S1. Stage 14L structured-feature reproducibility audit.",
        ROOT / "manuscript" / "figures" / "supp_figure_s1_stage14l_audit.png",
    ),
    (
        "Supplementary Figure S2. Optional post-hoc XAI example.",
        ROOT / "manuscript" / "figures" / "supp_figure_s2_xai_example.png",
    ),
]


def parse_bib(path: Path) -> dict[str, dict[str, str]]:
    text = path.read_text(encoding="utf-8")
    entries: dict[str, dict[str, str]] = {}
    for match in re.finditer(r"@\w+\{([^,]+),\s*(.*?)(?=\n@\w+\{|\Z)", text, re.S):
        key, body = match.group(1).strip(), match.group(2)
        fields: dict[str, str] = {}
        for fmatch in re.finditer(r"(\w+)\s*=\s*\{(.*?)\}\s*,?", body, re.S):
            fields[fmatch.group(1).lower()] = re.sub(r"\s+", " ", fmatch.group(2)).strip()
        entries[key] = fields
    return entries


def author_year(entry: dict[str, str]) -> str:
    authors = entry.get("author", "")
    year = entry.get("year", "n.d.")
    surnames = []
    for author in authors.split(" and "):
        author = author.strip()
        if not author:
            continue
        if "," in author:
            surnames.append(author.split(",", 1)[0].strip())
        else:
            surnames.append(author.split()[0].strip())
    if not surnames:
        return year
    if len(surnames) == 1:
        return f"{surnames[0]} {year}"
    if len(surnames) == 2:
        return f"{surnames[0]} and {surnames[1]} {year}"
    return f"{surnames[0]} et al. {year}"


def replace_citations(text: str, bib: dict[str, dict[str, str]]) -> str:
    def repl(match: re.Match[str]) -> str:
        raw = match.group(1)
        keys = [part.strip().lstrip("@") for part in raw.split(";")]
        labels = [author_year(bib.get(key, {})) for key in keys]
        return "(" + "; ".join(labels) + ")"

    return re.sub(r"\[@([^\]]+)\]", repl, text)


def format_reference(key: str, entry: dict[str, str]) -> str:
    authors = entry.get("author", "").replace(" and ", "; ")
    year = entry.get("year", "")
    title = entry.get("title", "")
    journal = entry.get("journal") or entry.get("booktitle") or entry.get("publisher", "")
    volume = entry.get("volume", "")
    pages = entry.get("pages", "")
    doi = entry.get("doi", "")
    parts = []
    if authors:
        parts.append(authors)
    if year:
        parts.append(f"({year})")
    if title:
        parts.append(title + ".")
    if journal:
        parts.append(journal + ".")
    detail = ", ".join([p for p in [volume, pages] if p])
    if detail:
        parts.append(detail + ".")
    if doi:
        parts.append(f"doi:{doi}")
    return " ".join(parts) if parts else key


def set_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)


def add_formatted_paragraph(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    # Basic Markdown bold support for manuscript fragments.
    pos = 0
    for match in re.finditer(r"\*\*(.*?)\*\*", text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    return paragraph


def add_bullets(doc: Document, line: str) -> None:
    add_formatted_paragraph(doc, line[2:].strip(), style="List Bullet")


def collect_used_keys(text: str) -> list[str]:
    used: list[str] = []
    for match in re.finditer(r"\[@([^\]]+)\]", text):
        for raw in match.group(1).split(";"):
            key = raw.strip().lstrip("@")
            if key and key not in used:
                used.append(key)
    return used


def build_docx(md_path: Path, out_path: Path, include_figures: bool = False) -> None:
    raw = md_path.read_text(encoding="utf-8")
    bib = parse_bib(BIB)
    used_keys = collect_used_keys(raw)
    text = replace_citations(raw, bib)

    doc = Document()
    set_style(doc)

    in_references_placeholder = False
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line == "## References":
            in_references_placeholder = True
            doc.add_heading("References", level=1)
            for i, key in enumerate(used_keys, 1):
                entry = bib.get(key, {})
                add_formatted_paragraph(doc, f"{i}. {format_reference(key, entry)}")
            continue
        if in_references_placeholder and line.startswith("References are managed"):
            continue
        if line.startswith("# "):
            p = add_formatted_paragraph(doc, line[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(16)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            add_bullets(doc, line)
            continue
        add_formatted_paragraph(doc, line)

    if include_figures:
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading("Embedded Figures for Review Copy", level=1)
        note = (
            "High-resolution PNG/PDF/SVG figure files are provided separately. "
            "The embedded images below are included for reviewer convenience."
        )
        add_formatted_paragraph(doc, note)
        for caption, path in FIGURES:
            if not path.exists():
                continue
            doc.add_heading(caption, level=2)
            doc.add_picture(str(path), width=Inches(6.2))

    doc.save(out_path)


if __name__ == "__main__":
    build_docx(MANUSCRIPT_MD, OUT_DOCX, include_figures=True)
    build_docx(COVER_MD, COVER_DOCX, include_figures=False)
    print(OUT_DOCX)
    print(COVER_DOCX)
