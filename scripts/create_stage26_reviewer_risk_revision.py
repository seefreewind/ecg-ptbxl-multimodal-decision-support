#!/usr/bin/env python3
"""Create Stage 26 reviewer-risk revision for the BMC MIDM ECG manuscript.

This is a conservative manuscript-only revision. It does not create new
experimental results; it clarifies wording, naming, and methods details that are
already supported by the code/configuration artifacts.
"""
from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
IN_DOC = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE25_CORE_FINDING_REVISED.docx"
OUT_DOC = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE26_REVIEWER_RISK_REVISED.docx"
LOG = ROOT / "manuscript" / "BMC_MIDM_STAGE26_REVIEWER_RISK_REVISION_LOG.md"
STAGE14L = ROOT / "tables" / "stage14l_internal_results.csv"
STAGE14L_OUT = ROOT / "manuscript" / "tables" / "supp_table_s10_reduced_schema_internal_threshold_diagnostics.csv"


def set_text(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def replace_in_paragraphs(doc: Document, replacements: dict[str, str]) -> None:
    for para in doc.paragraphs:
        txt = para.text
        new = txt
        for old, repl in replacements.items():
            new = new.replace(old, repl)
        if new != txt:
            set_text(para, new)


def replace_table_text(doc: Document, replacements: dict[str, str]) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = para.text
                    new = txt
                    for old, repl in replacements.items():
                        new = new.replace(old, repl)
                    if new != txt:
                        set_text(para, new)


def insert_after(doc: Document, needle: str, new_text: str) -> None:
    for i, para in enumerate(doc.paragraphs):
        if needle in para.text:
            para.insert_paragraph_before("__STAGE26_MARKER__")
            marker = doc.paragraphs[i]
            # Move marker text to after by swapping content with following paragraph is unreliable;
            # use XML insertion directly.
            p = para._p
            new_p = para._p.__class__()
            p.addnext(new_p)
            inserted = doc.paragraphs[i + 1]
            set_text(inserted, new_text)
            # Remove marker if present from previous attempt.
            if marker.text == "__STAGE26_MARKER__":
                marker._element.getparent().remove(marker._element)
            return
    raise RuntimeError(f"Needle not found: {needle}")


def para_after(doc: Document, needle: str, text: str) -> None:
    for para in doc.paragraphs:
        if needle in para.text:
            new_p = OxmlElement("w:p")
            para._p.addnext(new_p)
            inserted = None
            for p in doc.paragraphs:
                if p._p is new_p:
                    inserted = p
                    break
            if inserted is None:
                # python-docx may not refresh paragraph wrappers immediately; save/reload would work,
                # but constructing via _Paragraph is simpler and stable for this controlled edit.
                from docx.text.paragraph import Paragraph
                inserted = Paragraph(new_p, para._parent)
            inserted.add_run(text)
            return
    raise RuntimeError(f"Needle not found: {needle}")

def add_sensitivity_placeholder(doc: Document) -> None:
    text = (
        "The concordant-subset audit used a strict numerical agreement rule (absolute and relative tolerance 1e-6) "
        "to avoid treating implementation-dependent floating-point drift as interchangeable structured information. "
        "This threshold is intentionally conservative and may exclude features that are clinically similar but not numerically identical. "
        "The reduced-schema results should therefore be interpreted as a strict reproducibility audit rather than proof that no less stringent or de novo external feature engineering strategy could support multimodal validation. "
        "A tolerance-sensitivity analysis at looser numerical thresholds and a feature-count control analysis using randomly selected or importance-ranked 138-feature subsets were identified as useful follow-up checks, but were not used to revise the present test-set claims."
    )
    para_after(doc, "The decision rule was conservative.", text)


def add_f1_pipeline_note(doc: Document) -> None:
    note = (
        "Note. Gated-fusion F1 values refer to different evaluation pipelines: the main uncalibrated default-threshold test result is 0.7255 (Table 1), "
        "the temperature-scaled full-coverage diagnostic result is 0.7314, and the retained-80% uncertainty-triage subset is reported separately. "
        "These diagnostics should not replace the Table 1 primary model-comparison endpoint."
    )
    para_after(doc, "Gated fusion MLP achieved AUROC 0.9196, AP 0.7978, and F1 0.7255.", note)


def add_hyp_discussion(doc: Document) -> None:
    text = (
        "Per-superclass diagnostics suggested that much of the internal multimodal gain was concentrated in HYP: AUROC increased from 0.8359 to 0.8791, AP from 0.4839 to 0.5918, and F1 from 0.4444 to 0.5393 when matched concat was compared with the signal-embedding MLP. "
        "This pattern narrows the interpretation of modality complementarity: the released structured measurements appear most helpful for the hypertrophy superclass in this benchmark, while gains for NORM, MI, STTC, and CD were smaller and should be read descriptively."
    )
    para_after(doc, "The effect size was modest, however:", text)


def add_methods_details(doc: Document) -> None:
    for para in doc.paragraphs:
        txt = para.text
        if "The strong signal-only model used a one-dimensional residual network" in txt and "Adam-compatible" in txt:
            new = (
                "The strong signal-only model used a one-dimensional residual network with 12 input channels, 100 Hz waveforms, base channel width 64, three residual stages with two blocks per stage, dropout 0.2, and five output labels. Residual learning was used as the general architectural background [16]. The model provided both a direct waveform comparator and the 256-dimensional signal embedding for controlled multimodal comparison; embeddings were extracted from the penultimate representation of the trained signal model and then treated as frozen inputs for downstream MLP and matched-fusion experiments. Signal-embedding and structured MLP comparators used hidden dimensions 256/128 and 512/256, respectively, with dropout 0.3. The matched concat MLP used hidden dimensions 512/256, and the gated-fusion MLP used a 256-unit gating layer and a 256-unit classifier layer. Models used sigmoid multi-label outputs trained with binary cross-entropy with logits. Neural models used AdamW optimization in PyTorch with learning rate 0.001, weight decay 0.0001, batch size 64 for the signal model and 128 for MLP models, validation AUROC-based model selection, and early stopping patience of 8 epochs for the signal model and 10 epochs for MLP models. Maximum epochs were 50 for the signal model and 100 for MLP or gated models. The primary random seed was 2026; repeated-seed checks used seeds 2026, 2027, and 2028. The matched fusion dataset combined signal embeddings and released structured features after median imputation and standardization defined on the training set."
            )
            set_text(para, new)
            return
    raise RuntimeError("Could not find model-training details paragraph")

def add_waveform_preprocessing_details(doc: Document) -> None:
    text = (
        "Waveform records were read from the 100 Hz PTB-XL records100 WFDB files using wfdb.rdsamp. "
        "For signal-model inputs, each ECG record was represented as a 12-channel waveform and each lead was standardized within record to zero mean and unit variance with a small numerical floor on the standard deviation. "
        "The same signal-loader convention was used for signal-level external evaluation after dataset-specific WFDB compatibility checks. "
        "For the structured-feature reconstruction audit, candidate ECGdeli processing used the ECGdeli filtering and delineation workflow, including 1-40 Hz band-pass filtering, 50 Hz notch filtering, isoline correction, multichannel fiducial-point annotation, and subsequent amplitude or interval feature extraction where available."
    )
    para_after(doc, "The matched fusion dataset combined signal embeddings and released structured features after median imputation and standardization defined on the training set.", text)


def main() -> None:
    if not IN_DOC.exists():
        raise FileNotFoundError(IN_DOC)
    shutil.copy2(IN_DOC, OUT_DOC)
    doc = Document(OUT_DOC)

    replacements = {
        "The structured-feature reproducibility audit found 138 numerically concordant features, reduced-schema structured-only degradation, no stable reduced-schema multimodal gain, and only two joinable external structured records per dataset; external multimodal validation was therefore not supported.":
        "The structured-feature reproducibility audit found 138 numerically concordant features and reduced-schema internal degradation; external structured-feature coverage was far below any usable threshold, so external multimodal validation was not supported by the current extraction-and-joining pipeline.",
        "Risk stratification.": "Evidence boundary.",
        "Risk stratification": "Evidence boundary",
        "Fair comparison principle": "Matched-comparison principle",
        "fair model comparison": "matched model comparison",
        "fair fusion dataset": "matched fusion dataset",
        "fair multimodal comparison": "matched multimodal comparison",
        "fair fusion": "matched fusion",
        "fair-interface": "matched-interface",
        "fair MLP-concat": "matched MLP-concat",
        "fair concat": "matched concat",
        "Delta gated minus fair": "Delta gated minus matched concat",
        "gated minus fair": "gated minus matched concat",
        "stage14l_signal_embedding_mlp": "Signal-embedding MLP (reduced-schema audit)",
        "stage14l_structured_mlp": "Reduced structured-only MLP",
        "stage14l_fair_concat_mlp": "Reduced matched-concat MLP",
        "Table 6a Structured-feature reproducibility audit reduced-schema internal results": "Table 6a Structured-feature reproducibility audit reduced-schema internal results",
        "This framework supports reproducibility-aware evaluation of multimodal ECG diagnostic classification and demonstrates internal multimodal complementarity with signal-level external validation.":
        "This framework supports reproducibility-aware evaluation of multimodal ECG diagnostic classification. It showed modest internal multimodal complementarity, while external validation was limited to signal-only evaluation and did not test multimodal complementarity.",
        "current external structured-feature extraction and joining did not support external multimodal validation":
        "the current external structured-feature extraction-and-joining pipeline did not support external multimodal validation",
        "Structured-feature reproducibility as a central finding": "Structured-feature reproducibility as a feasibility boundary",
        "central finding": "feasibility boundary",
        "key contribution": "important audit finding",
    }
    replace_in_paragraphs(doc, replacements)
    replace_table_text(doc, replacements)

    add_methods_details(doc)
    add_waveform_preprocessing_details(doc)
    add_f1_pipeline_note(doc)
    add_hyp_discussion(doc)
    add_sensitivity_placeholder(doc)

    # Ensure headings and table lead-ins have consistent capitalization where they occur as plain paragraphs.
    for para in doc.paragraphs:
        if para.text.strip() == "structured-feature reproducibility audit":
            set_text(para, "Structured-Feature Reproducibility Audit")
        if para.text.strip() == "Table 3 Bootstrap comparison of gated fusion versus matched concat":
            set_text(para, "Table 3 Bootstrap Comparison of Gated Fusion Versus Matched Concat")

    doc.save(OUT_DOC)

    if STAGE14L.exists():
        df = pd.read_csv(STAGE14L)
        rename = {
            "stage14l_signal_embedding_mlp": "Signal-embedding MLP (reduced-schema audit)",
            "stage14l_structured_mlp": "Reduced structured-only MLP",
            "stage14l_fair_concat_mlp": "Reduced matched-concat MLP",
        }
        df["model"] = df["model"].replace(rename)
        STAGE14L_OUT.parent.mkdir(parents=True, exist_ok=True)
        df.to_csv(STAGE14L_OUT, index=False)

    LOG.write_text(
        "# BMC MIDM Stage 26 Reviewer-Risk Revision Log\n\n"
        "## Output\n"
        f"- Revised Word manuscript: `{OUT_DOC.relative_to(ROOT)}`\n"
        f"- This summary: `{LOG.relative_to(ROOT)}`\n"
        f"- Added threshold diagnostic table: `{STAGE14L_OUT.relative_to(ROOT)}`\n\n"
        "## Changes made\n"
        "- Softened abstract wording so the two-record external structured joinability result is not over-positioned as a biological or algorithmic proof.\n"
        "- Reframed the external structured-feature result as a current extraction-and-joining feasibility boundary, not proof that external PTB-XL+-compatible reconstruction is impossible in principle.\n"
        "- Added a strict-tolerance caveat for the 1e-6 concordance rule and explicitly named tolerance-sensitivity and 138-feature control analyses as future checks rather than completed results.\n"
        "- Clarified that reduced-schema degradation may reflect information loss from feature removal as well as reproducibility constraints.\n"
        "- Added an explicit gated-fusion F1 note distinguishing Table 1 default-threshold F1, temperature-scaled full-coverage F1, and retained-subset uncertainty F1.\n"
        "- Added a Discussion sentence noting that internal multimodal gain is concentrated most clearly in HYP.\n"
        "- Replaced internal code names in Table 6a with manuscript-readable model names.\n"
        "- Unified residual `fair` terminology toward `matched concat` / `matched comparison`.\n"
        "- Replaced `Risk stratification` keyword with `Evidence boundary`.\n"
        "- Added supported training details: sigmoid multi-label outputs, BCE-with-logits, AdamW, max epochs, early-stopping patience, and seeds.\n\n"
        "- Added waveform preprocessing details for WFDB loading, per-record lead standardization, and ECGdeli audit filtering/delineation steps.\n\n"
        "## Not done\n"
        "- No new ECGdeli external extraction was claimed.\n"
        "- No tolerance-sensitivity result was fabricated.\n"
        "- No random-138 or importance-138 feature-control result was fabricated.\n"
        "- No external multimodal validation claim was added.\n",
        encoding="utf-8",
    )

    print(OUT_DOC)
    print(LOG)
    if STAGE14L_OUT.exists():
        print(STAGE14L_OUT)


if __name__ == "__main__":
    main()
