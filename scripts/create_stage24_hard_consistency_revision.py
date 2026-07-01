#!/usr/bin/env python3
"""Create Stage 24 manuscript revision for hard consistency and methods detail."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript"
INPUT_DOCX = MANUSCRIPT_DIR / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE23_PRE_SUBMISSION_REVISED.docx"
OUTPUT_DOCX = MANUSCRIPT_DIR / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE24_HARDENED.docx"
LOG_MD = MANUSCRIPT_DIR / "BMC_MIDM_STAGE24_HARD_CONSISTENCY_REVISION_LOG.md"
TRIPOD_AI_CSV = MANUSCRIPT_DIR / "tables" / "supp_table_s7_tripod_ai_checklist.csv"


def replace_exact(doc: Document, old: str, new: str) -> bool:
    for p in doc.paragraphs:
        if p.text.strip() == old.strip():
            p.text = new
            return True
    return False


def replace_contains(doc: Document, old: str, new: str) -> int:
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
            count += 1
    return count


def find_para(doc: Document, startswith: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(startswith)


def insert_after(paragraph, text: str, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para.text = text
    if style is not None:
        new_para.style = style
    return new_para


def apply_spacing(doc: Document) -> None:
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 2
        p.paragraph_format.space_after = Pt(0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1
                    p.paragraph_format.space_after = Pt(0)


def normalize_abstract_label_bold(doc: Document) -> None:
    labels = ("Background:", "Methods:", "Results:", "Conclusions:")
    for p in doc.paragraphs:
        text = p.text
        for label in labels:
            if text.startswith(label):
                rest = text[len(label):]
                p.clear()
                r = p.add_run(label)
                r.bold = True
                p.add_run(rest)
                break


def remove_trailing_empty_paragraphs(doc: Document) -> int:
    removed = 0
    for p in reversed(doc.paragraphs):
        if p.text.strip():
            break
        element = p._element
        element.getparent().remove(element)
        removed += 1
    return removed


def main() -> None:
    doc = Document(INPUT_DOCX)
    changes: list[str] = []
    TRIPOD_AI_CSV.parent.mkdir(parents=True, exist_ok=True)

    # Title/framing: avoid over-promising risk stratification as prognosis.
    if replace_exact(
        doc,
        "A Reproducibility-Aware Decision-Support Evaluation Framework for Multimodal ECG Risk Stratification",
        "A Reproducibility-Aware Decision-Support Evaluation Framework for Multimodal ECG Diagnostic Classification",
    ):
        changes.append("Changed title from risk stratification to diagnostic classification.")

    # Abstract refinements.
    replace_contains(doc, "cardiac risk stratification", "cardiac diagnostic assessment")
    replace_contains(doc, "risk stratification", "diagnostic classification")
    replace_contains(doc, "CPSC2018 and Chapman-Shaoxing using pre-specified high-confidence label mappings", "CPSC2018 and the Chapman-Shaoxing-Ningbo ECG database using restricted high-confidence label subsets")
    replace_contains(doc, "on Chapman-Shaoxing, but Chapman-Shaoxing AP", "on the Chapman-Shaoxing-Ningbo high-confidence label subset, but AP")
    changes.append("Clarified external label subsets and replaced risk-stratification phrasing.")

    # Citation renumbering after adding PhysioNet ecg-arrhythmia citation.
    replace_contains(doc, "differences [13–15]", "differences [13–16]")
    replace_contains(doc, "CPSC2018 and Chapman-Shaoxing waveforms for signal-only evaluation [13,14]", "CPSC2018 and Chapman-Shaoxing-Ningbo waveforms for signal-only evaluation [13–15]")
    replace_contains(doc, "Residual learning was used as the general architectural background [16]", "Residual learning was used as the general architectural background [17]")
    replace_contains(doc, "medical AI cautions [9–10,17–18]", "medical AI cautions [9–10,18–19]")
    replace_contains(doc, "frozen internal test set [19]", "frozen internal test set [20]")
    changes.append("Inserted PhysioNet ecg-arrhythmia reference and shifted affected citation numbers.")

    # Data source and naming corrections.
    data_sources_old = (
        "CPSC2018 and Chapman-Shaoxing served as external waveform datasets for signal-only validation. "
        "CPSC2018 was evaluated for the high-confidence NORM/CD label subset. Chapman-Shaoxing was evaluated "
        "for MI/CD/HYP. Chapman-Shaoxing sinus rhythm was not treated as main-analysis NORM. Records that could "
        "not be read as WFDB waveforms were excluded. No external threshold tuning, external model selection, "
        "or external retraining was performed."
    )
    data_sources_new = (
        "CPSC2018 and the PhysioNet Chapman-Shaoxing-Ningbo ECG database served as external waveform datasets "
        "for signal-only validation. CPSC2018 was evaluated for the restricted high-confidence NORM/CD label "
        "subset. The Chapman-Shaoxing-Ningbo combined WFDB release was evaluated for the restricted MI/CD/HYP "
        "label subset; sinus rhythm labels were not treated as main-analysis NORM. The evaluated 45,150 records "
        "refer to this combined PhysioNet release after waveform-readability filtering, not to the earlier "
        "10,646-record Chapman-Shaoxing publication alone. Records that could not be read as WFDB waveforms "
        "were excluded. No external threshold tuning, external model selection, or external retraining was performed."
    )
    if replace_exact(doc, data_sources_old, data_sources_new):
        changes.append("Corrected Chapman-Shaoxing-Ningbo combined external dataset naming and sample-size context.")

    # Model/training details.
    model_para = find_para(doc, "The strong signal-only model used a one-dimensional residual network")
    model_para.text = (
        "The strong signal-only model used a one-dimensional residual network with 12 input channels, 100 Hz "
        "waveforms, base channel width 64, three residual stages with two blocks per stage, dropout 0.2, and "
        "five output labels. Residual learning was used as the general architectural background [17]. The model "
        "provided both a direct waveform comparator and the 256-dimensional signal embedding for controlled "
        "multimodal comparison. Signal-embedding and structured MLP comparators used hidden dimensions 256/128 "
        "and 512/256, respectively, with dropout 0.3. The fair concat MLP used hidden dimensions 512/256, and "
        "the gated-fusion MLP used a 256-unit gating layer and a 256-unit classifier layer. Neural models were "
        "trained with Adam-compatible optimization settings as implemented in PyTorch using learning rate 0.001, "
        "weight decay 0.0001, batch size 64 for the signal model and 128 for MLP models, validation AUROC-based "
        "model selection, and early stopping patience of 8 epochs for the signal model and 10 epochs for MLP models. "
        "The fair fusion dataset combined signal embeddings and released structured features after median imputation "
        "and standardization defined on the training set."
    )
    changes.append("Added concise architecture and training hyperparameter details.")

    # Structured-feature reproducibility criteria and join logic.
    audit_para = find_para(doc, "The structured-feature reproducibility analysis was treated")
    audit_para.text = (
        "The structured-feature reproducibility analysis was treated as a core feasibility audit. Instead of continuing "
        "attempts to reconstruct the full PTB-XL+ 531-column schema externally, the audit used a reduced structured-feature "
        "subset that had passed internal numerical concordance checks against released PTB-XL+ values. Candidate ECGdeli "
        "features were recomputed for a PTB-XL audit sample and compared with released PTB-XL+ values by feature name and "
        "ecg_id. A feature entered the reduced subset only when all non-missing paired values satisfied numpy-style "
        "allclose criteria with absolute tolerance 1e-6 and relative tolerance 1e-6; the manifest reports missingness, "
        "mean absolute error, median absolute error, and maximum absolute error for each selected feature. This subset "
        "contained 138 features. The reduced schema was tested internally with signal-embedding, structured-only, and "
        "fair concat models, and external coverage was checked for CPSC2018 and the Chapman-Shaoxing-Ningbo database."
    )
    insert_after(
        audit_para,
        "External joinability was defined operationally as the presence of both a signal-model prediction record and a candidate structured-feature row keyed by the standardized external record identifier after dataset-specific filename normalization. The two-record coverage result therefore supports a narrower conclusion: the current external extraction-and-joining pipeline did not provide enough matched structured rows for external multimodal evaluation. It should not be read as proof that compatible external structured-feature reconstruction is impossible in principle.",
        audit_para.style,
    )
    changes.append("Defined 138-feature numerical concordance criteria and external joinability boundary.")

    # Statistical and exclusions detail.
    stat_para = find_para(doc, "Discrimination was summarized using macro AUROC")
    stat_para.text = stat_para.text.replace(
        "The frozen PTB-XL split included 17,418 training records, 2,183 validation records, and 2,198 test records.",
        "The frozen PTB-XL split included 17,418 training records, 2,183 validation records, and 2,198 test records after records without the required five-superclass label representation or readable waveform/feature alignment were excluded from the analysis set.",
    )
    changes.append("Explained why the analyzed PTB-XL count differs from the full source-record count.")

    # Uncertainty F1 consistency + DCA threshold values.
    diag_para = find_para(doc, "Uncertainty triage, post-hoc explainability")
    diag_para.text = (
        "Uncertainty triage, post-hoc explainability, and exploratory decision-curve analysis were retained as "
        "decision-support diagnostics. In the gated-fusion uncertainty analysis, retaining the lowest-uncertainty "
        "80% of the test set retained 1,757 of 2,198 records and referred 441 records. In this retained subset, "
        "macro AUROC was 0.9303, AP was 0.8189, F1 was 0.7597, Brier score was 0.0687, and ECE was 0.0161. "
        "The corresponding full-coverage uncertainty-analysis values were macro AUROC 0.9196, AP 0.7978, "
        "temperature-scaled F1 0.7314, Brier score 0.0844, and ECE 0.0193; the primary Table 1 thresholded "
        "gated-fusion F1 remains 0.7255 under the main performance-evaluation pipeline. Case-level explainability "
        "summaries identified retained correct examples, retained errors, and high-uncertainty referred errors; "
        "top structured features and signal leads were reported as audit outputs rather than causal explanations. "
        "Exploratory internal decision-curve analysis is best interpreted from the threshold curves and representative "
        "thresholds rather than as a clinical decision rule. For fair concat, macro net benefit was 0.2146 at threshold "
        "0.10, 0.1908 at threshold 0.20, and 0.1724 at threshold 0.30; the corresponding net benefit versus treat-all "
        "was 0.0434, 0.1232, and 0.2380."
    )
    changes.append("Resolved gated F1 apparent mismatch and replaced DCA mean-only reporting with threshold-specific values.")

    # External results and discussion naming/label scope.
    replace_contains(doc, "Chapman-Shaoxing signal-only external validation included 45,150", "Chapman-Shaoxing-Ningbo signal-only external validation included 45,150")
    replace_contains(doc, "The Chapman-Shaoxing results", "The Chapman-Shaoxing-Ningbo results")
    replace_contains(doc, "Chapman-Shaoxing MI", "Chapman-Shaoxing-Ningbo MI")
    replace_contains(doc, "Chapman-Shaoxing had macro Brier", "Chapman-Shaoxing-Ningbo had macro Brier")
    replace_contains(doc, "Chapman-Shaoxing had 2 joinable", "Chapman-Shaoxing-Ningbo had 2 joinable")
    replace_contains(doc, "Chapman-Shaoxing showed preserved", "Chapman-Shaoxing-Ningbo showed preserved")
    replace_contains(doc, "Chapman-Shaoxing labels", "Chapman-Shaoxing-Ningbo labels")
    replace_contains(doc, "on Chapman-Shaoxing", "on Chapman-Shaoxing-Ningbo")
    replace_contains(doc, "for CPSC2018 and Chapman-Shaoxing", "for CPSC2018 and Chapman-Shaoxing-Ningbo")
    changes.append("Standardized external database naming across Results and Discussion.")

    # Limitations: multiple comparisons and external label scopes.
    limitations = find_para(doc, "Several limitations should be emphasized")
    limitations.text = limitations.text.replace(
        "Seventh, decision-curve analysis was exploratory, internal, and threshold-dependent.",
        "Seventh, bootstrap intervals and multiple diagnostic comparisons were reported descriptively without multiplicity correction. Eighth, decision-curve analysis was exploratory, internal, and threshold-dependent.",
    ).replace(
        "Eighth, no clinical workflow",
        "Ninth, no clinical workflow",
    )
    changes.append("Added multiplicity limitation.")

    # Supplementary/reporting checklist note.
    supp_list = find_para(doc, "Supplementary Table S6:")
    insert_after(supp_list, "Supplementary Table S7: TRIPOD+AI reporting checklist.", supp_list.style)
    changes.append("Added TRIPOD+AI checklist as expected supplementary material.")

    # Declarations: move AI note mention from Methods to a declaration-compatible phrase without deleting Methods detail.
    replace_exact(
        doc,
        "AI-assisted tools",
        "AI-assisted tools and author responsibility",
    )

    # Data availability naming and DOI.
    data_avail = find_para(doc, "PTB-XL is available from PhysioNet")
    data_avail.text = (
        "PTB-XL is available from PhysioNet at https://physionet.org/content/ptb-xl/. PTB-XL+ is available "
        "from PhysioNet at https://physionet.org/content/ptb-xl-plus/. CPSC2018 is available from the China "
        "Physiological Signal Challenge 2018 website at http://2018.icbeb.org/Challenge.html. The external "
        "Chapman-Shaoxing-Ningbo ECG database used here is the PhysioNet large-scale 12-lead ECG database for "
        "arrhythmia study at https://physionet.org/content/ecg-arrhythmia/ with DOI https://doi.org/10.13026/wgex-er52. "
        "The analysis code and manuscript artifacts are available at https://github.com/seefreewind/ecg-ptbxl-multimodal-decision-support "
        "and archived on Zenodo at https://doi.org/10.5281/zenodo.21091784."
    )
    changes.append("Updated data-availability wording with the PhysioNet ecg-arrhythmia DOI.")

    # Reference list corrections and insertion.
    replace_exact(
        doc,
        "15. Alday EAP, Gu A, Shah AJ, Robichaux C, Wong AKI, Liu C, Liu F, Rad AB, Elola A, Seyedi S, et al.. Classification of 12-lead ECGs: the PhysioNet/Computing in Cardiology Challenge 2020. Physiological Measurement. 2020;41:124003. doi:10.1088/1361-6579/abc960.",
        "16. Alday EAP, Gu A, Shah AJ, Robichaux C, Wong AKI, Liu C, Liu F, Rad AB, Elola A, Seyedi S, et al. Classification of 12-lead ECGs: the PhysioNet/Computing in Cardiology Challenge 2020. Physiological Measurement. 2020;41:124003. doi:10.1088/1361-6579/abc960.",
    )
    replace_exact(
        doc,
        "16. He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016:770–778.",
        "17. He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016:770–778.",
    )
    replace_exact(
        doc,
        "17. Selvaraju RR, Cogswell M, Das A, Vedantam R, Parikh D, Batra D. Grad-CAM: Visual explanations from deep networks via gradient-based localization. IEEE International Conference on Computer Vision. 2017:618–626.",
        "18. Selvaraju RR, Cogswell M, Das A, Vedantam R, Parikh D, Batra D. Grad-CAM: Visual explanations from deep networks via gradient-based localization. IEEE International Conference on Computer Vision. 2017:618–626.",
    )
    replace_exact(
        doc,
        "18. Ghassemi M, Oakden-Rayner L, Beam AL. The false hope of current approaches to explainable artificial intelligence in health care. The Lancet Digital Health. 2021;3:e745–e750. doi:10.1016/S2589-7500(21)00208-9.",
        "19. Ghassemi M, Oakden-Rayner L, Beam AL. The false hope of current approaches to explainable artificial intelligence in health care. The Lancet Digital Health. 2021;3:e745–e750. doi:10.1016/S2589-7500(21)00208-9.",
    )
    replace_exact(
        doc,
        "19. Efron B, Tibshirani RJ. An Introduction to the Bootstrap. Chapman and Hall/CRC. 1993.",
        "20. Efron B, Tibshirani RJ. An Introduction to the Bootstrap. Chapman and Hall/CRC. 1993.",
    )
    ref14 = find_para(doc, "14. Zheng J, Zhang J")
    insert_after(
        ref14,
        "15. Zheng J, Guo H, Chu H. A large scale 12-lead electrocardiogram database for arrhythmia study. PhysioNet. 2022. doi:10.13026/wgex-er52.",
        ref14.style,
    )
    changes.append("Added PhysioNet ecg-arrhythmia dataset citation and fixed et al. punctuation.")

    # Figure/supplement titles capitalization and naming.
    replace_contains(doc, "Supplementary Figure S1. structured-feature reproducibility audit", "Supplementary Figure S1. Structured-feature reproducibility audit")
    replace_contains(doc, "Table 6a structured-feature reproducibility audit", "Table 6a Structured-feature reproducibility audit")
    replace_contains(doc, "Table 6b structured-feature reproducibility audit", "Table 6b Structured-feature reproducibility audit")
    replace_contains(doc, "the not supported decision", "the unsupported decision")
    replace_contains(doc, "Chapman-Shaoxing per-class", "Chapman-Shaoxing-Ningbo per-class")
    replace_contains(doc, "Chapman-Shaoxing MI/CD/HYP", "Chapman-Shaoxing-Ningbo MI/CD/HYP")

    trailing_removed = remove_trailing_empty_paragraphs(doc)
    if trailing_removed:
        changes.append(f"Removed {trailing_removed} trailing blank paragraph(s) after embedded figure previews.")

    normalize_abstract_label_bold(doc)
    apply_spacing(doc)
    doc.save(OUTPUT_DOCX)

    tripod_rows = [
        ["domain", "item", "manuscript_location", "status", "note"],
        ["Title/Abstract", "Identify study as public-data model development/evaluation", "Title; Abstract", "addressed", "Diagnostic classification framing added in Stage 24."],
        ["Background", "State rationale and intended model role", "Background", "addressed", "Decision-support evaluation and evidence-boundary role stated."],
        ["Data", "Describe data sources and eligibility", "Methods: Data sources and roles", "addressed", "PTB-XL/PTB-XL+, CPSC2018, and Chapman-Shaoxing-Ningbo roles specified."],
        ["Participants/records", "Report sample sizes and exclusions", "Methods: Statistical analysis; Data sources", "addressed", "PTB-XL split sizes and external readable-record counts reported."],
        ["Predictors/features", "Describe signal and structured inputs", "Methods: Model groups; Structured-feature reproducibility audit", "addressed", "Signal embeddings and released PTB-XL+ structured features described."],
        ["Outcome/labels", "Define target labels and label scope", "Methods: Data sources and roles", "addressed", "Internal five superclasses and restricted external label subsets stated."],
        ["Model development", "Describe architecture and training", "Methods: Model groups", "addressed", "Architecture and key hyperparameters summarized."],
        ["Validation", "Describe internal and external validation", "Methods; Results", "addressed", "Internal multimodal and signal-only external validation separated."],
        ["Performance measures", "Report discrimination, calibration, and thresholded metrics", "Methods: Statistical analysis; Results", "addressed", "AUROC, AP, F1, Brier, ECE, MCE, bootstrap CIs reported."],
        ["Uncertainty/calibration", "State calibration and uncertainty methods", "Methods; Results", "addressed", "Temperature scaling and uncertainty triage described."],
        ["Explainability", "Describe XAI limits", "Methods; Results; Discussion", "addressed", "XAI treated as audit output, not causal explanation."],
        ["Clinical use", "Avoid unsupported deployment claims", "Discussion; Limitations", "addressed", "No clinical-readiness claim made."],
        ["Reproducibility", "Provide code/data availability", "Declarations", "addressed", "GitHub and Zenodo DOI included."],
        ["External multimodal limitation", "State unsupported external multimodal validation", "Results; Discussion", "addressed", "External evidence remains signal-only."],
    ]
    with TRIPOD_AI_CSV.open("w", newline="", encoding="utf-8") as fh:
        csv.writer(fh).writerows(tripod_rows)

    log = [
        "# BMC MIDM Stage 24 Hard Consistency Revision Log",
        "",
        "## Purpose",
        "",
        "This revision addresses hard consistency, external-dataset naming, and methods-detail issues identified in the latest pre-submission read.",
        "",
        "## Changes",
        "",
    ]
    log.extend(f"- {c}" for c in changes)
    log.extend(
        [
            "",
            "## Preserved Boundaries",
            "",
            "- External evidence remains signal-only.",
            "- The structured-feature audit is reported as an extraction/joinability and concordance feasibility audit, not as proof that external feature reconstruction is impossible.",
            "- No gated-fusion superiority or clinical-readiness claim was added.",
            "",
            "## Generated Files",
            "",
            f"- `{OUTPUT_DOCX.relative_to(ROOT)}`",
            f"- `{LOG_MD.relative_to(ROOT)}`",
            f"- `{TRIPOD_AI_CSV.relative_to(ROOT)}`",
        ]
    )
    LOG_MD.write_text("\n".join(log) + "\n", encoding="utf-8")
    print(OUTPUT_DOCX)
    print(LOG_MD)


if __name__ == "__main__":
    main()
