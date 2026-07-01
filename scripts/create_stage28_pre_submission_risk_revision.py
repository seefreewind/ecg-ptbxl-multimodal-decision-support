#!/usr/bin/env python3
"""Create Stage 28 pre-submission risk-reduction manuscript revision.

This revision is prose/format focused. It does not add new analyses and does
not change the locked evidence boundary: external validation remains
signal-only, and external multimodal validation remains unsupported.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
IN_DOC = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE27_ROBUSTNESS_REVISED.docx"
OUT_DOC = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE28_PRE_SUBMISSION_REVISED.docx"
LOG = ROOT / "manuscript" / "BMC_MIDM_STAGE28_PRE_SUBMISSION_RISK_REVISION_LOG.md"


def set_text(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].bold = False
    else:
        paragraph.add_run(text)


def set_labelled(paragraph, label: str, body: str) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    first = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    first.text = label
    first.bold = True
    second = paragraph.add_run(" " + body)
    second.bold = False


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def para_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.add_run(text)
    return inserted


def replace_contains(doc: Document, needle: str, replacement: str) -> None:
    for para in doc.paragraphs:
        if needle in para.text:
            set_text(para, replacement)
            return
    raise RuntimeError(f"Could not find paragraph containing: {needle}")


def delete_contains(doc: Document, needle: str) -> None:
    for para in doc.paragraphs:
        if needle in para.text:
            delete_paragraph(para)
            return


def round_metric_text(text: str) -> str:
    mapping = {
        "0.9193": "0.919",
        "0.7953": "0.795",
        "0.7208": "0.721",
        "0.9098": "0.910",
        "0.7721": "0.772",
        "0.6998": "0.700",
        "0.9094": "0.909",
        "0.7724": "0.772",
        "0.7002": "0.700",
        "0.9046": "0.905",
        "0.7652": "0.765",
        "0.6899": "0.690",
        "0.9196": "0.920",
        "0.7978": "0.798",
        "0.7255": "0.726",
        "0.9071": "0.907",
        "0.6509": "0.651",
        "0.5904": "0.590",
        "0.8742": "0.874",
        "0.1727": "0.173",
        "0.1650": "0.165",
        "0.5704": "0.570",
        "0.3045": "0.304",
        "0.0000": "0.000",
        "0.9097": "0.910",
        "0.7731": "0.773",
        "0.6938": "0.694",
        "0.9195": "0.920",
        "0.7960": "0.796",
        "0.7110": "0.711",
        "0.9221": "0.922",
        "0.8031": "0.803",
        "0.7310": "0.731",
        "0.0027": "0.003",
        "0.0796": "0.080",
        "0.0277": "0.028",
        "0.0677": "0.068",
        "0.2624": "0.262",
        "0.3319": "0.332",
        "0.0166": "0.017",
        "0.1760": "0.176",
        "0.1353": "0.135",
        "0.000201": "0.00020",
        "0.000044": "0.00004",
    }
    for old, new in mapping.items():
        text = text.replace(old, new)
    text = text.replace("+0.0098", "+0.010")
    text = text.replace("+0.0229", "+0.023")
    text = text.replace("+0.0205", "+0.021")
    text = text.replace("+0.0094", "+0.009")
    text = text.replace("+0.0232", "+0.023")
    text = text.replace("+0.0209", "+0.021")
    text = text.replace("+0.0003", "+0.000")
    text = text.replace("+0.0025", "+0.003")
    text = text.replace("+0.0047", "+0.005")
    return text


def round_docx_tables(doc: Document) -> None:
    numeric = re.compile(r"^[+-]?\d+\.\d{4,}$")
    signed = re.compile(r"^[+-]\d+\.\d{4,}$")
    ci = re.compile(r"([+-]?\d+\.\d{4,})")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if numeric.match(text) or signed.match(text):
                        value = float(text)
                        prefix = "+" if text.startswith("+") and value >= 0 else ""
                        set_text(para, f"{prefix}{value:.3f}")
                    elif " to " in text and ci.search(text):
                        def repl(match):
                            value = float(match.group(1))
                            prefix = "+" if match.group(1).startswith("+") and value >= 0 else ""
                            return f"{prefix}{value:.3f}"
                        set_text(para, ci.sub(repl, text))


def main() -> None:
    shutil.copy2(IN_DOC, OUT_DOC)
    doc = Document(OUT_DOC)

    set_labelled(
        doc.paragraphs[8],
        "Results:",
        (
            "With the released full PTB-XL+ schema, matched multimodal concatenation achieved AUROC 0.919, average precision 0.795, and F1 0.721, exceeding the signal-embedding comparator by +0.010 AUROC, +0.023 average precision, and +0.021 F1. "
            "Gated fusion had slightly higher point estimates but no statistically clear advantage over concatenation. The structured-feature reproducibility audit found 138 numerically concordant features, but this concordance-selected subset did not preserve internal multimodal gain. "
            "Robustness controls showed that random and internally importance-ranked 138-feature subsets recovered internal gains, indicating loss of informative structured features rather than a feature-count limit alone. "
            "External multimodal validation was not attempted because the current external structured-feature extraction-and-joining workflow produced candidate structured rows for only two records per external dataset. "
            "Signal-only external validation achieved macro AUROC 0.907 on CPSC2018 and 0.874 on the Chapman-Shaoxing-Ningbo high-confidence label subset, with low AP/F1 in low-prevalence labels."
        ),
    )
    set_labelled(
        doc.paragraphs[9],
        "Conclusions:",
        (
            "Released PTB-XL+ features supported modest internal multimodal complementarity, but current external structured-feature extraction was not sufficiently complete to support external multimodal validation. "
            "The most informative robustness finding was that numerically reproducible features were not necessarily the features carrying multimodal signal. "
            "Externally reproducible, sufficiently informative, and complete structured-feature extraction is needed before external multimodal or clinical-use claims can be considered."
        ),
    )

    replace_contains(
        doc,
        "The gap addressed here is therefore not only model performance.",
        (
            "The gap addressed here is therefore not only model performance. A central methodological question is whether the features that are easiest to reproduce numerically are also the features that carry clinically and statistically useful multimodal signal. "
            "Our robustness analysis showed that this need not be true: concordance-selected features were reproducible but weak, whereas random and internally importance-ranked 138-feature controls recovered internal gains. "
            "This distinction gives the study a positive methodological focus: responsible multimodal ECG evaluation must test not only internal performance, but also whether externally reconstructed structured features are both reproducible and informative."
        ),
    )

    replace_contains(
        doc,
        "The concordant-subset audit used a strict numerical agreement rule",
        (
            "The concordant-subset audit used a strict numerical agreement rule (absolute and relative tolerance 1e-6) to avoid treating implementation-dependent floating-point drift as interchangeable structured information. "
            "A robustness audit evaluated tolerance sensitivity and 138-feature count controls. The feature count remained 138 at 1e-4 and 1e-3 and increased only to 141 at 1e-2; substantially looser tolerances admitted more features but were not treated as exact reproduction. "
            "The same audit compared the concordance-selected 138 features with random and internally attribution-ranked 138-feature subsets under the same frozen internal split and validation-only model-selection rules. "
            "These controls were used to interpret the failure mechanism of the concordance-selected subset, not to establish external multimodal validation."
        ),
    )

    replace_contains(
        doc,
        "Decision-curve analysis was evaluated as an exploratory",
        (
            "Decision-curve analysis was evaluated as an exploratory, threshold-dependent analysis [11]. "
            "Label-specific net-benefit curves were used as supplementary decision-support diagnostics. "
            "Aggregated net-benefit summaries were not treated as standard clinical-utility endpoints and were not used to claim readiness for practice."
        ),
    )
    replace_contains(
        doc,
        "Exploratory decision-curve analysis was calculated per label",
        (
            "Exploratory decision-curve analysis was calculated per label and reported as supplementary source data. "
            "Because net benefit is conventionally defined for a specific decision threshold and outcome, cross-label aggregate summaries were treated only as descriptive audit outputs and were not emphasized as main clinical-utility endpoints."
        ),
    )

    delete_contains(doc, "Gated-fusion F1 values refer to different evaluation pipelines")
    replace_contains(
        doc,
        "Matched MLP-concat achieved AUROC",
        (
            "Matched MLP-concat achieved AUROC 0.919, AP 0.795, and F1 0.721. Gated fusion MLP achieved AUROC 0.920, AP 0.798, and F1 0.726. "
            "Although the gated model had slightly higher raw point estimates, paired bootstrap intervals for gated minus matched concat contained zero; the simpler matched concat model was therefore used as the primary multimodal comparator."
        ),
    )
    replace_contains(
        doc,
        "Uncertainty triage, post-hoc explainability, and exploratory decision-curve analysis",
        (
            "Uncertainty triage, post-hoc explainability, and exploratory decision-curve analysis were retained as decision-support diagnostics and are reported primarily in supplementary materials. "
            "In the gated-fusion uncertainty analysis, retaining the lowest-uncertainty 80% of the test set retained 1,757 of 2,198 records and referred 441 records, with more favorable discrimination and calibration among retained cases. "
            "These analyses were used to audit model behavior, not to replace the main uncalibrated Table 1 default-threshold performance endpoints."
        ),
    )

    replace_contains(
        doc,
        "Chapman-Shaoxing-Ningbo signal-only external validation included",
        (
            "Chapman-Shaoxing-Ningbo signal-only external validation included 45,150 evaluated records for the MI/CD/HYP high-confidence label scope. Macro AUROC was 0.874, macro AP was 0.173, and macro F1 was 0.165. "
            "These results should be interpreted mainly as ranking behavior under label shift rather than as strong thresholded-performance evidence. "
            "The MI mapping was especially weak for clinical interpretation: only 123 records were mapped as MI-positive (prevalence 0.003), with AP 0.080 and F1 0.028. "
            "This very low prevalence suggests that the Chapman-Shaoxing-Ningbo MI label scope is not directly comparable with the PTB-XL MI superclass; it is retained as a label-mapping stress test rather than as a robust external MI performance result."
        ),
    )

    replace_contains(
        doc,
        "External structured-feature coverage also failed the quality gate.",
        (
            "External structured-feature reconstruction did not reach a scale suitable for evaluation. In the current external audit, candidate ECGdeli-derived structured rows were available for only 2 CPSC2018 records among 9,944 signal prediction records and only 2 Chapman-Shaoxing-Ningbo records among 45,150 signal prediction records. "
            "Because the candidate feature records were themselves sparse, this result is best interpreted as an engineering feasibility limitation of the present extraction-and-joining workflow, not as a scientific estimate that PTB-XL+-compatible external reconstruction is impossible in principle."
        ),
    )
    delete_contains(doc, "The audit did not support external multimodal validation. This should be presented")

    replace_contains(
        doc,
        "This study evaluated a reproducibility-aware decision-support framework",
        (
            "This study evaluated a reproducibility-aware decision-support framework for multimodal ECG diagnostic classification. "
            "The principal finding is not that external multimodal validation failed as a model comparison, because such validation was not performed. "
            "Rather, the study shows a reproducibility-informativeness gap: released PTB-XL+ features supported modest internal multimodal complementarity, but the numerically concordant features that could be reproduced most strictly were not the features that preserved internal multimodal gain, and current external feature extraction produced too few candidate rows for multimodal testing."
        ),
    )
    replace_contains(
        doc,
        "The internal full-schema results showed",
        (
            "The internal full-schema results showed that released PTB-XL+ structured features added value beyond ECG signal embeddings under frozen PTB-XL splits. "
            "The signal model used here should be viewed as a controlled comparator rather than a leaderboard-optimized PTB-XL classifier; stronger published PTB-XL benchmarks may report higher signal-only AUROC with larger architectures or training recipes. "
            "This design choice was deliberate because the study's target was matched multimodal evidence attribution, calibration, and reproducibility boundaries rather than state-of-the-art waveform classification."
        ),
    )
    replace_contains(
        doc,
        "The external evidence is intentionally limited",
        (
            "The external evidence is intentionally limited to signal-only validation. CPSC2018 supported limited signal-level external transportability for the NORM/CD high-confidence label subset. "
            "Chapman-Shaoxing-Ningbo showed preserved ranking performance but weak AP and F1, especially in low-prevalence labels. "
            "The MI result should be interpreted as a label-shift stress test rather than robust diagnostic evidence because only 123 of 45,150 evaluated records were MI-positive under the mapping used here."
        ),
    )
    replace_contains(
        doc,
        "The structured-feature reproducibility audit is an important",
        (
            "The structured-feature reproducibility audit is an important BMC MIDM-compatible contribution because it separates internal reuse of released PTB-XL+ features from de novo reconstruction of comparable external structured features. "
            "The robustness controls sharpened this interpretation. Random and internally importance-ranked 138-feature subsets recovered internal multimodal gains, whereas the concordance-selected 138-feature subset did not. "
            "The audit should therefore be read as evidence of a reproducibility-informativeness gap: the easiest features to reproduce numerically were not the features carrying most of the internal multimodal signal. "
            "The practical implication remains conservative: released internal PTB-XL+ values can support internal multimodal experiments, but the present external WFDB reconstruction workflow is not complete enough for external multimodal validation."
        ),
    )

    replace_contains(
        doc,
        "This study used publicly available, de-identified datasets.",
        (
            "This study used only publicly available, de-identified datasets. Ethical approval and consent procedures were handled by the original dataset creators as described in the corresponding dataset publications and repositories. "
            "No new human participants were recruited and no identifiable private health information was accessed. Under the authors' institutional policies for secondary analysis of public de-identified data, no additional institutional review board approval was required."
        ),
    )

    # Remove internal stage numbering from manuscript-facing supplementary titles.
    replacements = {
        "Stage 27 robustness audit": "Robustness audit",
        "A Stage 27 robustness audit": "A robustness audit",
        "The Stage 27 controls": "The robustness controls",
        "Stage 27 controls": "Robustness controls",
        "Stage 27 tolerance-sensitivity audit": "Tolerance-sensitivity audit",
        "Stage 27 feature-count control results": "Feature-count control results",
        "Stage 27 external join-failure audit": "External join-failure audit",
        "Evidence boundary.": "Evidence boundary.",
    }
    for para in doc.paragraphs:
        new = para.text
        for old, repl in replacements.items():
            new = new.replace(old, repl)
        new = round_metric_text(new)
        if new != para.text:
            # Preserve abstract labels if present.
            if new.startswith("Results:"):
                set_labelled(para, "Results:", new[len("Results:"):].strip())
            elif new.startswith("Conclusions:"):
                set_labelled(para, "Conclusions:", new[len("Conclusions:"):].strip())
            else:
                set_text(para, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    new = para.text
                    for old, repl in replacements.items():
                        new = new.replace(old, repl)
                    new = round_metric_text(new)
                    if new != para.text:
                        set_text(para, new)
    round_docx_tables(doc)
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 4 and cells[0].text.strip() == "CPSC2018" and cells[3].text.strip() == "0.000":
                set_text(cells[3].paragraphs[0], "0.00020")
            if len(cells) >= 4 and cells[0].text.strip() == "Chapman-Shaoxing-Ningbo" and cells[3].text.strip() == "0.000":
                set_text(cells[3].paragraphs[0], "0.00004")

    doc.save(OUT_DOC)

    LOG.write_text(
        "# BMC MIDM Stage 28 Pre-Submission Risk Revision Log\n\n"
        "## Output\n"
        f"- Revised Word manuscript: `{OUT_DOC.relative_to(ROOT)}`\n"
        f"- This summary: `{LOG.relative_to(ROOT)}`\n\n"
        "## Main changes\n"
        "- Reframed external structured-feature coverage as incomplete candidate extraction/joining in the current workflow, not a scientific proof of impossibility.\n"
        "- Foregrounded the reproducibility-informativeness gap: concordance-selected features were reproducible but weak, whereas random/importance 138-feature controls recovered internal gains.\n"
        "- Removed manuscript-facing `Stage 27` terminology and replaced it with descriptive robustness-audit language.\n"
        "- Simplified gated-fusion F1 presentation by keeping Table 1 as the primary endpoint and moving uncertainty/calibration F1 variants into supplementary-diagnostic framing.\n"
        "- Downweighted macro DCA language and treated DCA as supplementary label-specific decision-support diagnostics.\n"
        "- Strengthened Chapman-Shaoxing-Ningbo MI caveat as a label-shift stress test rather than robust external MI evidence.\n"
        "- Added a controlled-comparator explanation for why the signal model is not positioned as PTB-XL SOTA.\n"
        "- Rewrote ethics wording for public de-identified secondary analysis.\n"
        "- Rounded manuscript-facing performance values to three decimals where practical; machine-readable source tables retain full precision.\n\n"
        "## Not changed\n"
        "- No external multimodal claim was added.\n"
        "- No exact PTB-XL+ external reconstruction claim was added.\n"
        "- No gated-fusion superiority claim was added.\n",
        encoding="utf-8",
    )
    print(OUT_DOC)
    print(LOG)


if __name__ == "__main__":
    main()
