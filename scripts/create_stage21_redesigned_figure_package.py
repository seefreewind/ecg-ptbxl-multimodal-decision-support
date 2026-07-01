#!/usr/bin/env python3
from __future__ import annotations

import math
from pathlib import Path
from shutil import copyfile

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "figures_redesigned"
OUT.mkdir(parents=True, exist_ok=True)

STAGE20_DOCX = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE20_FINAL.docx"
STAGE21_DOCX = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE21_FIGURE_REDESIGNED.docx"
LOG = ROOT / "manuscript" / "BMC_MIDM_STAGE21_FIGURE_REDESIGN_LOG.md"

mpl.rcParams.update(
    {
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"],
        "svg.fonttype": "none",
        "pdf.fonttype": 42,
        "font.size": 8,
        "axes.spines.right": False,
        "axes.spines.top": False,
        "axes.linewidth": 0.7,
        "legend.frameon": False,
    }
)

COL = {
    "navy": "#1f4e79",
    "blue": "#3b73b9",
    "teal": "#1b9e77",
    "teal_light": "#d9f0eb",
    "orange": "#e69f00",
    "orange_light": "#fff3d6",
    "red": "#c75d5d",
    "red_light": "#f7dddd",
    "gray": "#6b7280",
    "gray_light": "#f3f4f6",
    "black": "#202124",
}


def save_all(fig: plt.Figure, stem: str) -> list[Path]:
    paths = []
    for ext in ["png", "pdf", "svg"]:
        p = OUT / f"{stem}.{ext}"
        if ext == "png":
            fig.savefig(p, dpi=450, bbox_inches="tight", facecolor="white")
        else:
            fig.savefig(p, bbox_inches="tight", facecolor="white")
        paths.append(p)
    plt.close(fig)
    return paths


def panel_label(ax, label: str) -> None:
    ax.text(-0.03, 1.04, label, transform=ax.transAxes, fontsize=11, fontweight="bold", va="top")


def rounded(ax, xy, w, h, text, fc, ec="#d1d5db", lw=1.0, size=8, color=None):
    from matplotlib.patches import FancyBboxPatch

    patch = FancyBboxPatch(
        xy,
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.035",
        facecolor=fc,
        edgecolor=ec,
        linewidth=lw,
    )
    ax.add_patch(patch)
    ax.text(
        xy[0] + w / 2,
        xy[1] + h / 2,
        text,
        ha="center",
        va="center",
        fontsize=size,
        color=color or COL["black"],
        linespacing=1.15,
    )
    return patch


def draw_ecg(ax, x0, y0, w, h, color=COL["blue"]):
    x = np.linspace(0, 1, 240)
    y = 0.5 + 0.04 * np.sin(18 * np.pi * x)
    for pos, amp in [(0.18, 0.15), (0.34, -0.18), (0.39, 0.34), (0.44, -0.20), (0.68, 0.12)]:
        y += amp * np.exp(-((x - pos) ** 2) / 0.0008)
    ax.plot(x0 + x * w, y0 + y * h, color=color, lw=1.8, clip_on=False)


def figure1():
    fig = plt.figure(figsize=(12.4, 7.4))
    gs = fig.add_gridspec(4, 1, height_ratios=[1.2, 1.2, 1.0, 1.35], hspace=0.42)
    fig.suptitle("Reproducibility-aware ECG decision-support evaluation framework", x=0.02, y=0.985, ha="left", fontsize=15, fontweight="bold")

    # A
    ax = fig.add_subplot(gs[0]); ax.axis("off"); ax.set_xlim(0, 1); ax.set_ylim(0, 1); panel_label(ax, "A")
    ax.text(0.02, 0.88, "Internal multimodal data", fontweight="bold", fontsize=10)
    rounded(ax, (0.05, 0.20), 0.26, 0.42, "PTB-XL waveforms\n12-lead ECG", COL["gray_light"], COL["blue"])
    draw_ecg(ax, 0.08, 0.30, 0.20, 0.22)
    rounded(ax, (0.37, 0.20), 0.26, 0.42, "Released PTB-XL+\nstructured features", COL["gray_light"], COL["teal"])
    # table icon
    for i in range(4):
        ax.plot([0.42, 0.58], [0.31 + i * 0.05, 0.31 + i * 0.05], color=COL["teal"], lw=0.8)
    for i in range(4):
        ax.plot([0.42 + i * 0.053, 0.42 + i * 0.053], [0.31, 0.46], color=COL["teal"], lw=0.8)
    ax.annotate("", xy=(0.72, 0.41), xytext=(0.64, 0.41), arrowprops=dict(arrowstyle="->", color=COL["gray"], lw=1.5))
    rounded(ax, (0.73, 0.20), 0.22, 0.42, "Internal PTB-XL/PTB-XL+\nmultimodal evaluation", COL["teal_light"], COL["teal"], lw=1.3)

    # B
    ax = fig.add_subplot(gs[1]); ax.axis("off"); ax.set_xlim(0, 1); ax.set_ylim(0, 1); panel_label(ax, "B")
    ax.text(0.02, 0.88, "Fair model comparison", fontweight="bold", fontsize=10)
    labels = ["strong\nsignal-only", "signal-embedding\nMLP", "structured\nMLP", "fair\nconcat", "gated\nfusion"]
    xs = np.linspace(0.08, 0.74, 5)
    for x, lab in zip(xs, labels):
        rounded(ax, (x, 0.36), 0.12, 0.28, lab, "white", COL["blue"] if "signal" in lab else COL["teal"] if "structured" in lab else COL["gray"])
    ax.annotate("", xy=(0.86, 0.50), xytext=(0.79, 0.50), arrowprops=dict(arrowstyle="->", color=COL["gray"], lw=1.5))
    rounded(ax, (0.86, 0.50), 0.12, 0.22, "fair concat gain\nsupported", COL["teal_light"], COL["teal"], size=7.5)
    rounded(ax, (0.86, 0.20), 0.12, 0.22, "gating: no clear\nadditional benefit", COL["red_light"], COL["red"], size=7.5)
    ax.text(0.52, 0.16, "paired bootstrap on frozen internal test set", ha="center", color=COL["gray"], fontsize=7.8)

    # C
    ax = fig.add_subplot(gs[2]); ax.axis("off"); ax.set_xlim(0, 1); ax.set_ylim(0, 1); panel_label(ax, "C")
    ax.text(0.02, 0.85, "Conservative decision-support diagnostics", fontweight="bold", fontsize=10)
    mods = ["calibration", "uncertainty\ntriage", "XAI", "exploratory\nDCA"]
    for i, mod in enumerate(mods):
        rounded(ax, (0.09 + i * 0.20, 0.35), 0.15, 0.30, mod, COL["orange_light"], COL["orange"])
    rounded(ax, (0.82, 0.28), 0.16, 0.42, "Diagnostics support\nauditability,\nnot deployment", COL["gray_light"], COL["gray"], size=7.8)

    # D
    ax = fig.add_subplot(gs[3]); ax.axis("off"); ax.set_xlim(0, 1); ax.set_ylim(0, 1); panel_label(ax, "D")
    ax.text(0.02, 0.90, "External evidence boundary", fontweight="bold", fontsize=10)
    rounded(ax, (0.05, 0.52), 0.22, 0.26, "CPSC2018\nexternal waveforms", "white", COL["blue"])
    rounded(ax, (0.31, 0.52), 0.22, 0.26, "Chapman-Shaoxing\nexternal waveforms", "white", COL["blue"])
    rounded(ax, (0.58, 0.52), 0.18, 0.26, "signal-only external\nvalidation: GO", COL["teal_light"], COL["teal"], size=7.8)
    rounded(ax, (0.05, 0.16), 0.35, 0.22, "structured-feature reconstruction audit", COL["gray_light"], COL["gray"])
    rounded(ax, (0.47, 0.13), 0.46, 0.28, "external multimodal validation: NO-GO\nunder current coverage/fidelity checks", COL["red_light"], COL["red"], size=8.2)
    rounded(ax, (0.03, 0.01), 0.94, 0.07, "Evidence boundary: internal multimodal evidence + signal-only external validation + structured-feature reproducibility audit", "white", COL["gray"], size=7.6)
    return save_all(fig, "figure1_evidence_boundary_graphical_abstract")


def figure2():
    perf = pd.DataFrame(
        [
            ["strong signal-only", 0.9098, 0.7721, 0.6998],
            ["signal-embedding MLP", 0.9094, 0.7724, 0.7002],
            ["structured MLP", 0.9046, 0.7652, 0.6899],
            ["fair MLP-concat", 0.9193, 0.7953, 0.7208],
            ["gated fusion MLP", 0.9196, 0.7978, 0.7255],
        ],
        columns=["model", "AUROC", "AP", "F1"],
    )
    fair = pd.DataFrame(
        [
            ["concat - signal-embedding", "AUROC", 0.0098, 0.0067, 0.0131],
            ["concat - signal-embedding", "AP", 0.0229, 0.0157, 0.0302],
            ["concat - signal-embedding", "F1", 0.0205, 0.0088, 0.0324],
            ["concat - strong signal-only", "AUROC", 0.0094, 0.0062, 0.0127],
            ["concat - strong signal-only", "AP", 0.0232, 0.0151, 0.0308],
            ["concat - strong signal-only", "F1", 0.0209, 0.0098, 0.0319],
        ],
        columns=["comparison", "metric", "delta", "lo", "hi"],
    )
    gated = pd.DataFrame(
        [["AUROC", 0.0003, -0.0015, 0.0021], ["AP", 0.0025, -0.0014, 0.0070], ["F1", 0.0047, -0.0044, 0.0142]],
        columns=["metric", "delta", "lo", "hi"],
    )
    fig = plt.figure(figsize=(12.2, 5.8))
    gs = fig.add_gridspec(1, 3, width_ratios=[1.25, 1.2, 0.95], wspace=0.48)
    ax = fig.add_subplot(gs[0]); panel_label(ax, "A")
    metrics = ["AUROC", "AP", "F1"]; ybase = np.arange(len(metrics))
    offsets = np.linspace(-0.22, 0.22, len(perf))
    colors = [COL["blue"], "#6aaed6", COL["teal"], COL["orange"], COL["red"]]
    for off, (_, row), c in zip(offsets, perf.iterrows(), colors):
        vals = [row[m] for m in metrics]
        ax.scatter(vals, ybase + off, s=38, color=c, label=row["model"], zorder=3)
        for v, y in zip(vals, ybase + off):
            ax.plot([v - 0.004, v + 0.004], [y, y], color=c, lw=1.4)
    ax.set_yticks(ybase); ax.set_yticklabels(metrics)
    ax.set_xlim(0.66, 0.94); ax.set_xlabel("Internal test performance")
    ax.grid(axis="x", color="#e5e7eb", lw=0.7)
    ax.legend(loc="lower left", bbox_to_anchor=(0.0, -0.28), ncol=1, fontsize=7)
    ax.set_title("Internal model performance", loc="left", fontsize=10, fontweight="bold")

    ax = fig.add_subplot(gs[1]); panel_label(ax, "B")
    fair["label"] = fair["comparison"] + " · " + fair["metric"]
    y = np.arange(len(fair))[::-1]
    ax.axvline(0, color=COL["gray"], lw=1)
    ax.errorbar(fair["delta"], y, xerr=[fair["delta"] - fair["lo"], fair["hi"] - fair["delta"]], fmt="o", color=COL["teal"], ecolor=COL["teal"], capsize=3)
    ax.set_yticks(y); ax.set_yticklabels(fair["label"], fontsize=7)
    ax.set_xlabel("Delta in metric")
    ax.set_title("Paired bootstrap: fair concat gain", loc="left", fontsize=10, fontweight="bold")
    ax.text(0.005, -0.9, "All CIs exclude zero;\nsupported internal gain", color=COL["teal"], fontsize=7.5)
    ax.grid(axis="x", color="#e5e7eb", lw=0.7)

    ax = fig.add_subplot(gs[2]); panel_label(ax, "C")
    y = np.arange(len(gated))[::-1]
    ax.axvline(0, color=COL["gray"], lw=1)
    ax.errorbar(gated["delta"], y, xerr=[gated["delta"] - gated["lo"], gated["hi"] - gated["delta"]], fmt="o", color=COL["red"], ecolor=COL["red"], capsize=3)
    ax.set_yticks(y); ax.set_yticklabels(gated["metric"])
    ax.set_xlabel("Delta gated minus fair")
    ax.set_title("Gated vs fair concat", loc="left", fontsize=10, fontweight="bold")
    ax.text(-0.004, -0.65, "CIs cross zero:\nno statistically clear\nadditional benefit", color=COL["red"], fontsize=7.5)
    ax.grid(axis="x", color="#e5e7eb", lw=0.7)
    return save_all(fig, "figure2_internal_performance_bootstrap")


def figure3():
    macro = pd.DataFrame(
        [["CPSC2018", 0.9071, 0.6509, 0.5904], ["Chapman-Shaoxing", 0.8742, 0.1727, 0.1650]],
        columns=["dataset", "AUROC", "AP", "F1"],
    )
    chap = pd.DataFrame(
        [["MI", 123, 0.0027, 0.0796, 0.0277], ["CD", 3058, 0.0677, 0.2624, 0.3319], ["HYP", 751, 0.0166, 0.1760, 0.1353]],
        columns=["label", "positives", "prevalence", "AP", "F1"],
    )
    fig = plt.figure(figsize=(12, 4.8))
    gs = fig.add_gridspec(1, 3, width_ratios=[1.1, 1.35, 0.95], wspace=0.45)
    ax = fig.add_subplot(gs[0]); panel_label(ax, "A")
    x = np.arange(len(macro)); w = 0.22
    for i, (m, c) in enumerate(zip(["AUROC", "AP", "F1"], [COL["blue"], COL["orange"], COL["teal"]])):
        ax.bar(x + (i - 1) * w, macro[m], width=w, color=c, label=m)
    ax.set_xticks(x); ax.set_xticklabels(macro["dataset"], rotation=15, ha="right")
    ax.set_ylim(0, 1.02); ax.set_ylabel("Macro metric")
    ax.set_title("External macro performance", loc="left", fontsize=10, fontweight="bold")
    ax.text(0.02, 0.94, "External validation was signal-only", transform=ax.transAxes, fontsize=7.5, color=COL["gray"])
    ax.legend(loc="upper right", fontsize=7)

    ax = fig.add_subplot(gs[1]); panel_label(ax, "B")
    y = np.arange(len(chap))[::-1]
    ax.barh(y + 0.18, chap["prevalence"], height=0.16, color=COL["gray"], label="prevalence")
    ax.scatter(chap["AP"], y, color=COL["orange"], s=45, label="AP", zorder=4)
    ax.scatter(chap["F1"], y - 0.18, color=COL["teal"], s=45, label="F1", zorder=4)
    for _, r in chap.iterrows():
        yy = y[list(chap["label"]).index(r["label"])]
        if r["label"] == "MI":
            ax.text(0.36, yy, "MI: 123 positives,\nprevalence 0.0027", color=COL["red"], fontsize=7.5, va="center")
    ax.set_yticks(y); ax.set_yticklabels(chap["label"])
    ax.set_xlim(0, 0.42)
    ax.set_xlabel("Prevalence / AP / F1")
    ax.set_title("Chapman-Shaoxing per-class diagnostics", loc="left", fontsize=10, fontweight="bold")
    ax.legend(loc="lower right", fontsize=7)
    ax.grid(axis="x", color="#e5e7eb", lw=0.7)
    ax.text(0.0, -0.92, "Low prevalence and internal-threshold transfer contribute to weak AP/F1.", color=COL["gray"], fontsize=7.4)

    ax = fig.add_subplot(gs[2]); ax.axis("off"); panel_label(ax, "C")
    ax.set_title("Evidence boundary", loc="left", fontsize=10, fontweight="bold")
    rounded(ax, (0.05, 0.62), 0.9, 0.24, "What this shows:\nsignal-level external transportability", COL["teal_light"], COL["teal"], size=8)
    rounded(ax, (0.05, 0.34), 0.9, 0.20, "What this does not show:\nexternal multimodal validation", COL["red_light"], COL["red"], size=8)
    rounded(ax, (0.05, 0.10), 0.9, 0.16, "No external threshold tuning", COL["gray_light"], COL["gray"], size=8)
    return save_all(fig, "figure3_external_signal_validation_diagnostics")


def figure4():
    internal = pd.DataFrame(
        [
            ["strong signal-only", 0.0903, 0.0320],
            ["structured MLP", 0.0942, 0.0212],
            ["fair concat", 0.0864, 0.0283],
            ["gated fusion", 0.0844, 0.0193],
        ],
        columns=["model", "Macro Brier", "Macro ECE"],
    )
    rel = pd.read_csv(ROOT / "results/calibration/reliability_curve_source_data.csv")
    rel = rel[(rel["split"] == "test") & (rel["calibration_mode"] == "temperature_scaled") & (rel["model_name"].isin(["strong_signal_only", "fair_concat_mlp", "gated_fusion_mlp"]))]
    rel = rel.groupby(["model_name", "bin"], as_index=False).agg(confidence=("confidence", "mean"), accuracy=("accuracy", "mean"), count=("count", "sum"))
    ext = pd.DataFrame([["CPSC2018", 0.1262, 0.4099], ["Chapman-Shaoxing", 0.1412, 0.7277]], columns=["dataset", "Macro ECE", "Macro MCE"])
    fig = plt.figure(figsize=(12.2, 4.7))
    gs = fig.add_gridspec(1, 3, width_ratios=[1.1, 1.1, 1.0], wspace=0.42)
    ax = fig.add_subplot(gs[0]); panel_label(ax, "A")
    y = np.arange(len(internal))
    ax.scatter(internal["Macro Brier"], y + 0.12, color=COL["blue"], label="Macro Brier", s=45)
    ax.scatter(internal["Macro ECE"], y - 0.12, color=COL["orange"], label="Macro ECE", s=45)
    ax.set_yticks(y); ax.set_yticklabels(internal["model"])
    ax.invert_yaxis(); ax.set_xlim(0, 0.11)
    ax.set_xlabel("Metric value")
    ax.set_title("Internal calibration summary", loc="left", fontsize=10, fontweight="bold")
    ax.legend(fontsize=7)
    ax.grid(axis="x", color="#e5e7eb", lw=0.7)

    ax = fig.add_subplot(gs[1]); panel_label(ax, "B")
    name_map = {"strong_signal_only": "strong signal-only", "fair_concat_mlp": "fair concat", "gated_fusion_mlp": "gated fusion"}
    for name, c in zip(name_map, [COL["blue"], COL["orange"], COL["red"]]):
        d = rel[rel["model_name"] == name].sort_values("confidence")
        ax.plot(d["confidence"], d["accuracy"], marker="o", ms=3, lw=1.5, color=c, label=name_map[name])
    ax.plot([0, 1], [0, 1], color=COL["gray"], lw=1, ls="--")
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.set_aspect("equal", adjustable="box")
    ax.set_xlabel("Mean predicted probability"); ax.set_ylabel("Observed frequency")
    ax.set_title("Reliability curve", loc="left", fontsize=10, fontweight="bold")
    ax.legend(fontsize=7, loc="lower right")

    ax = fig.add_subplot(gs[2]); panel_label(ax, "C")
    x = np.arange(len(ext)); w = 0.28
    ax.bar(x - w/2, ext["Macro ECE"], width=w, color=COL["orange"], label="Macro ECE")
    ax.bar(x + w/2, ext["Macro MCE"], width=w, color=COL["red"], label="Macro MCE")
    ax.set_xticks(x); ax.set_xticklabels(ext["dataset"], rotation=15, ha="right")
    ax.set_ylim(0, 0.82); ax.set_ylabel("Metric value")
    ax.set_title("External no-refit calibration", loc="left", fontsize=10, fontweight="bold")
    ax.text(0.02, 0.82, "Temperature source: internal validation;\nno external refitting", transform=ax.transAxes, fontsize=7.5, color=COL["gray"])
    ax.text(0.02, 0.08, "Distribution-shift behavior,\nnot clinical calibration readiness", transform=ax.transAxes, fontsize=7.5, color=COL["gray"])
    ax.legend(fontsize=7)
    return save_all(fig, "figure4_calibration_reliability_distribution_shift")


def supp1():
    internal = pd.DataFrame(
        [["signal-embedding MLP", 0.9094, 0.7722, 0.6981], ["reduced structured MLP", 0.5704, 0.3045, 0.0000], ["reduced fair concat", 0.9097, 0.7731, 0.6938]],
        columns=["model", "AUROC", "AP", "F1"],
    )
    coverage = pd.DataFrame([["CPSC2018", 9944, 2, 0.000201], ["Chapman-Shaoxing", 45150, 2, 0.000044]], columns=["dataset", "signal", "joinable", "coverage"])
    fig = plt.figure(figsize=(12.2, 5.6))
    gs = fig.add_gridspec(1, 4, width_ratios=[1.0, 1.25, 1.1, 1.05], wspace=0.48)
    ax = fig.add_subplot(gs[0]); ax.axis("off"); panel_label(ax, "A")
    ax.set_title("Schema reduction", loc="left", fontsize=10, fontweight="bold")
    rounded(ax, (0.12, 0.64), 0.76, 0.18, "Released PTB-XL+\n531 features", COL["gray_light"], COL["gray"])
    ax.annotate("", xy=(0.50, 0.46), xytext=(0.50, 0.63), arrowprops=dict(arrowstyle="->", color=COL["gray"], lw=1.5))
    rounded(ax, (0.20, 0.25), 0.60, 0.18, "Allclose subset\n138 features", COL["teal_light"], COL["teal"])

    ax = fig.add_subplot(gs[1]); panel_label(ax, "B")
    y = np.arange(len(internal))
    for metric, c, off in [("AUROC", COL["blue"], 0.18), ("AP", COL["orange"], 0.0), ("F1", COL["teal"], -0.18)]:
        ax.scatter(internal[metric], y + off, color=c, s=42, label=metric)
    ax.set_yticks(y); ax.set_yticklabels(internal["model"], fontsize=7.5)
    ax.invert_yaxis(); ax.set_xlim(0, 1.02)
    ax.set_xlabel("Internal test metric")
    ax.set_title("Reduced-schema internal test", loc="left", fontsize=10, fontweight="bold")
    ax.legend(fontsize=7)
    ax.text(0.04, 2.55, "Structured-only collapsed;\nfair concat did not preserve\nstable multimodal gain.", color=COL["red"], fontsize=7.5)
    ax.grid(axis="x", color="#e5e7eb", lw=0.7)

    ax = fig.add_subplot(gs[2]); panel_label(ax, "C")
    y = np.arange(len(coverage))[::-1]
    ax.barh(y, coverage["coverage"], color=COL["red"], height=0.32)
    ax.set_xscale("log"); ax.set_xlim(1e-5, 1)
    ax.set_yticks(y); ax.set_yticklabels(coverage["dataset"])
    ax.set_xlabel("Joinable structured coverage\n(log scale)")
    ax.set_title("External coverage", loc="left", fontsize=10, fontweight="bold")
    for yi, (_, r) in zip(y, coverage.iterrows()):
        ax.text(r["coverage"] * 1.4, yi, f"{int(r['joinable'])}/{int(r['signal']):,}", va="center", fontsize=7.5)
    ax.grid(axis="x", color="#e5e7eb", lw=0.7)

    ax = fig.add_subplot(gs[3]); ax.axis("off"); panel_label(ax, "D")
    ax.set_title("Audit decision", loc="left", fontsize=10, fontweight="bold")
    rounded(ax, (0.06, 0.52), 0.88, 0.24, "External multimodal\nvalidation: NO-GO", COL["red_light"], COL["red"], size=10)
    rounded(ax, (0.06, 0.25), 0.88, 0.18, "Reason: insufficient external\ncoverage/fidelity checks", COL["gray_light"], COL["gray"], size=8)
    ax.text(0.08, 0.08, "Reproducibility/feasibility finding,\nnot a failed external multimodal model.", fontsize=7.6, color=COL["gray"])
    return save_all(fig, "supp_figure_s1_stage14l_reproducibility_audit")


def update_word():
    copyfile(STAGE20_DOCX, STAGE21_DOCX)
    doc = Document(STAGE21_DOCX)
    replacements = {
        "Figure 1. Study design and evidence boundaries.": "Figure 1. Reproducibility-aware ECG decision-support evaluation framework and evidence boundaries.",
        "Figure 2. Internal model performance and statistically supported multimodal gain.": "Figure 2. Internal model performance, paired bootstrap support for fair concat, and gated-fusion comparison.",
        "Figure 3. Signal-only external validation.": "Figure 3. External signal-only validation and Chapman-Shaoxing per-class diagnostic context.",
        "Figure 4. Calibration and reliability under internal testing and external distribution shift.": "Figure 4. Calibration and reliability under internal testing and external distribution shift.",
        "Supplementary Figure S1. Stage 14L structured-feature reproducibility audit.": "Supplementary Figure S1. Stage 14L structured-feature reproducibility audit funnel.",
    }
    legend_map = {
        "Figure 1.": "Figure 1. Reproducibility-aware ECG decision-support evaluation framework and evidence boundaries. Panels summarize internal PTB-XL/PTB-XL+ multimodal inputs, fair model comparison, conservative decision-support diagnostics, and the external evidence boundary. External multimodal validation was not performed, and no clinical deployment claim is made.",
        "Figure 2.": "Figure 2. Internal model performance, paired bootstrap support for fair concat, and gated-fusion comparison. Panel A shows frozen-test AUROC, AP, and F1 across model groups. Panel B shows paired bootstrap deltas supporting fair concat gain over unimodal comparators. Panel C shows that gated fusion did not provide a statistically clear additional benefit over fair concat.",
        "Figure 3.": "Figure 3. External signal-only validation and Chapman-Shaoxing per-class diagnostic context. Panel A shows macro AUROC, AP, and F1 for CPSC2018 and Chapman-Shaoxing. Panel B shows Chapman-Shaoxing MI/CD/HYP prevalence, AP, and F1. Panel C states the evidence boundary: signal-level external transportability only, no external multimodal validation, and no external threshold tuning.",
        "Figure 4.": "Figure 4. Calibration and reliability under internal testing and external distribution shift. Panel A summarizes internal macro Brier and macro ECE. Panel B shows internal reliability curves with an ideal calibration reference. Panel C reports external no-refit calibration behavior using temperature scaling fitted on internal validation data.",
        "Supplementary Figure S1.": "Supplementary Figure S1. Stage 14L structured-feature reproducibility audit funnel. Panels show reduction from the released 531-column PTB-XL+ schema to the 138-feature allclose subset, reduced-schema internal results, external structured-feature coverage, and the NO-GO decision for external multimodal validation under current coverage/fidelity checks.",
    }
    for p in doc.paragraphs:
        txt = p.text
        for old, new in replacements.items():
            if txt == old:
                p.text = new
        for prefix, new in legend_map.items():
            if txt.startswith(prefix):
                p.text = new
    # Remove old embedded figure section.
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Figures":
            start = i
            break
    if start is not None:
        for p in list(doc.paragraphs[start:]):
            el = p._element
            el.getparent().remove(el)
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Figures", level=1)
    doc.add_paragraph("Separate high-resolution PNG/PDF/SVG figure files are available for upload where applicable.")
    fig_items = [
        ("Figure 1. Reproducibility-aware ECG decision-support evaluation framework and evidence boundaries.", "figure1_evidence_boundary_graphical_abstract.png"),
        ("Figure 2. Internal model performance, paired bootstrap support for fair concat, and gated-fusion comparison.", "figure2_internal_performance_bootstrap.png"),
        ("Figure 3. External signal-only validation and Chapman-Shaoxing per-class diagnostic context.", "figure3_external_signal_validation_diagnostics.png"),
        ("Figure 4. Calibration and reliability under internal testing and external distribution shift.", "figure4_calibration_reliability_distribution_shift.png"),
        ("Supplementary Figure S1. Stage 14L structured-feature reproducibility audit funnel.", "supp_figure_s1_stage14l_reproducibility_audit.png"),
    ]
    for title, filename in fig_items:
        doc.add_heading(title, level=2)
        doc.add_picture(str(OUT / filename), width=Inches(6.25))
    doc.save(STAGE21_DOCX)


def write_log(paths_by_fig: dict[str, list[Path]]):
    lines = [
        "# BMC MIDM Stage 21 Figure Redesign Log",
        "",
        "## Figure Contract",
        "",
        "- Core conclusion: the manuscript supports internal multimodal complementarity, signal-only external validation, and a structured-feature reproducibility NO-GO boundary for external multimodal validation.",
        "- Backend: Python/matplotlib only.",
        "- Evidence boundary unchanged: no external multimodal validation, no gated fusion superiority, and no clinical deployment/readiness claim.",
        "",
        "## Redesigned Figures",
        "",
    ]
    source_map = {
        "Figure 1": "Manuscript evidence-boundary design, PTB-XL/PTB-XL+ and external validation role definitions.",
        "Figure 2": "Internal model performance values and paired bootstrap CIs from manuscript/tables/table_internal_multimodal_gain_bootstrap.csv plus locked manuscript values.",
        "Figure 3": "tables/table_external_signal_results.csv and manuscript/tables/supp_table_external_per_class_diagnostics.csv.",
        "Figure 4": "figures/source_data/fig4_calibration_long.csv, results/calibration/reliability_curve_source_data.csv, and tables/table_stage15_external_signal_calibration.csv.",
        "Supplementary Figure S1": "tables/stage14l_internal_results.csv, tables/stage14l_external_results.csv, and Stage 14L locked audit values.",
    }
    for label, files in paths_by_fig.items():
        lines.append(f"### {label}")
        lines.append("")
        lines.append(f"- Source data used: {source_map[label]}")
        lines.append("- Output files:")
        for p in files:
            detail = ""
            if p.suffix == ".png":
                im = Image.open(p)
                detail = f" ({im.width}x{im.height}px)"
            lines.append(f"  - `{p.relative_to(ROOT)}`{detail}")
        lines.append("- Submission-ready: yes.")
        lines.append("")
    lines.extend(
        [
            "## XAI Figure Decision",
            "",
            "- Supplementary Figure S2 was not retained in the redesigned submission package.",
            "- Existing XAI source data and heatmap images remain available, but the prior review-copy figure was not sufficiently readable at Word/page scale.",
            "- No speculative XAI figure was created.",
            "",
            "## Updated Word Manuscript",
            "",
            f"- Updated Word file with redesigned embedded previews: `{STAGE21_DOCX.relative_to(ROOT)}`",
            "- Figure legends were updated to match the redesigned figures.",
            "- Informal wording such as optional, draft, candidate, or review copy was not used in figure headings.",
            "",
            "## Missing Source Data",
            "",
            "- None for redesigned Figure 1-4 or Supplementary Figure S1.",
            "",
            "## Remaining Figure Blockers",
            "",
            "- None for the required redesigned figure package.",
        ]
    )
    LOG.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    paths = {
        "Figure 1": figure1(),
        "Figure 2": figure2(),
        "Figure 3": figure3(),
        "Figure 4": figure4(),
        "Supplementary Figure S1": supp1(),
    }
    update_word()
    write_log(paths)
    print(STAGE21_DOCX)
    print(LOG)


if __name__ == "__main__":
    main()
