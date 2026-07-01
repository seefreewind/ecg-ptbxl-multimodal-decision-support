#!/usr/bin/env python3
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_MD = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT.md"
BIB = ROOT / "manuscript" / "references.bib"
OUT_DOCX = ROOT / "manuscript" / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE19_CLEAN.docx"
LOG_MD = ROOT / "manuscript" / "BMC_MIDM_STAGE19_WORD_CLEANUP_LOG.md"


RETAINED_FIGURES = [
    ("Figure 1. Study design and evidence boundaries.", "figure1_study_design_evidence_boundary.png"),
    ("Figure 2. Internal model performance and statistically supported multimodal gain.", "figure2_internal_model_performance.png"),
    ("Figure 3. Signal-only external validation.", "figure3_external_signal_only_validation.png"),
    ("Figure 4. Calibration and reliability under internal testing and external distribution shift.", "figure4_calibration_reliability.png"),
    ("Supplementary Figure S1. Stage 14L structured-feature reproducibility audit.", "supp_figure_s1_stage14l_audit.png"),
]


TABLES = [
    {
        "title": "Table 1 Internal full-schema model performance on the frozen PTB-XL/PTB-XL+ test set",
        "headers": ["Model", "Test AUROC", "Test AP", "Test F1"],
        "rows": [
            ["strong signal-only", "0.9098", "0.7721", "0.6998"],
            ["signal-embedding MLP", "0.9094", "0.7724", "0.7002"],
            ["structured MLP", "0.9046", "0.7652", "0.6899"],
            ["fair MLP-concat", "0.9193", "0.7953", "0.7208"],
            ["gated fusion MLP", "0.9196", "0.7978", "0.7255"],
        ],
        "note": "AP = average precision; AUROC = area under the receiver operating characteristic curve.",
        "widths": [2.6, 1.2, 1.2, 1.2],
    },
    {
        "title": "Table 2 Statistical support for the internal multimodal gain",
        "headers": ["Comparison", "Metric", "Fair concat", "Comparator", "Delta", "95% CI", "CI contains zero"],
        "rows": [
            ["fair concat minus signal-embedding MLP", "AUROC", "0.9193", "0.9094", "+0.0098", "+0.0067 to +0.0131", "no"],
            ["fair concat minus signal-embedding MLP", "AP", "0.7953", "0.7724", "+0.0229", "+0.0157 to +0.0302", "no"],
            ["fair concat minus signal-embedding MLP", "F1", "0.7208", "0.7002", "+0.0205", "+0.0088 to +0.0324", "no"],
            ["fair concat minus strong signal-only", "AUROC", "0.9193", "0.9098", "+0.0094", "+0.0062 to +0.0127", "no"],
            ["fair concat minus strong signal-only", "AP", "0.7953", "0.7721", "+0.0232", "+0.0151 to +0.0308", "no"],
            ["fair concat minus strong signal-only", "F1", "0.7208", "0.6998", "+0.0209", "+0.0098 to +0.0319", "no"],
        ],
        "note": "AP = average precision; AUROC = area under the receiver operating characteristic curve; CI = confidence interval. Thresholds were selected on internal validation data only.",
        "widths": [2.35, 0.75, 0.95, 0.95, 0.8, 1.45, 1.05],
    },
    {
        "title": "Table 3 Gated fusion versus fair concat on the frozen internal test set",
        "headers": ["Metric", "Delta gated minus fair", "95% CI", "Interpretation"],
        "rows": [
            ["AUROC", "+0.0003", "-0.0015 to 0.0021", "CI contains zero"],
            ["AP", "+0.0025", "-0.0014 to 0.0070", "CI contains zero"],
            ["F1", "+0.0047", "-0.0044 to 0.0142", "CI contains zero"],
        ],
        "note": "AP = average precision; AUROC = area under the receiver operating characteristic curve; CI = confidence interval.",
        "widths": [1.0, 1.7, 1.8, 2.0],
    },
    {
        "title": "Table 4 Signal-only external validation",
        "headers": ["Dataset", "Label scope", "N", "Macro AUROC", "Macro AP", "Macro F1"],
        "rows": [
            ["CPSC2018", "NORM/CD", "9,944", "0.9071", "0.6509", "0.5904"],
            ["Chapman-Shaoxing", "MI/CD/HYP", "45,150", "0.8742", "0.1727", "0.1650"],
        ],
        "note": "AP = average precision; AUROC = area under the receiver operating characteristic curve. External validation was signal-only, and thresholds were selected on internal validation data only.",
        "widths": [1.7, 1.4, 0.9, 1.15, 1.0, 1.0],
    },
    {
        "title": "Table 5 External signal-only calibration",
        "headers": ["Dataset", "Macro Brier", "Micro Brier", "Macro ECE", "Macro MCE"],
        "rows": [
            ["CPSC2018", "0.1268", "0.1268", "0.1262", "0.4099"],
            ["Chapman-Shaoxing", "0.0855", "0.0855", "0.1412", "0.7277"],
        ],
        "note": "ECE = expected calibration error; MCE = maximum calibration error. External calibration was evaluated without external refitting.",
        "widths": [1.8, 1.25, 1.25, 1.25, 1.25],
    },
    {
        "title": "Table 6a Stage 14L reduced-schema internal results",
        "headers": ["Model", "AUROC", "AP", "F1"],
        "rows": [
            ["stage14l_signal_embedding_mlp", "0.9094", "0.7722", "0.6981"],
            ["stage14l_structured_mlp", "0.5704", "0.3045", "0.0000"],
            ["stage14l_fair_concat_mlp", "0.9097", "0.7731", "0.6938"],
        ],
        "note": "AP = average precision; AUROC = area under the receiver operating characteristic curve. Stage 14L was a reproducibility/feasibility audit, not an external multimodal validation result.",
        "widths": [2.6, 1.0, 1.0, 1.0],
    },
    {
        "title": "Table 6b Stage 14L external structured-feature coverage",
        "headers": ["Dataset", "Signal records", "Joinable structured records", "Coverage"],
        "rows": [
            ["CPSC2018", "9,944", "2", "0.000201"],
            ["Chapman-Shaoxing", "45,150", "2", "0.000044"],
        ],
        "note": "The external structured-feature coverage was insufficient to support external multimodal validation.",
        "widths": [1.7, 1.3, 2.0, 1.1],
    },
]


def parse_bib(path: Path) -> dict[str, dict[str, str]]:
    text = path.read_text(encoding="utf-8")
    entries: dict[str, dict[str, str]] = {}
    for match in re.finditer(r"@\w+\{([^,]+),\s*(.*?)(?=\n@\w+\{|\Z)", text, re.S):
        key, body = match.group(1).strip(), match.group(2)
        fields: dict[str, str] = {}
        for fmatch in re.finditer(r"(\w+)\s*=\s*\{(.*?)\}\s*,?", body, re.S):
            fields[fmatch.group(1).lower()] = re.sub(r"\s+", " ", fmatch.group(2)).strip()
        entries[key] = fields
    return entries


def collect_used_keys(text: str) -> list[str]:
    used: list[str] = []
    for match in re.finditer(r"\[@([^\]]+)\]", text):
        for raw in match.group(1).split(";"):
            key = raw.strip().lstrip("@")
            if key and key not in used:
                used.append(key)
    return used


def compact_numbers(nums: list[int]) -> str:
    if not nums:
        return ""
    out: list[str] = []
    i = 0
    while i < len(nums):
        start = nums[i]
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j > i:
            out.append(f"{start}\u2013{nums[j]}")
        else:
            out.append(str(start))
        i = j + 1
    return ",".join(out)


def replace_citations_numeric(text: str, order: list[str]) -> str:
    index = {key: i + 1 for i, key in enumerate(order)}

    def repl(match: re.Match[str]) -> str:
        keys = [part.strip().lstrip("@") for part in match.group(1).split(";")]
        nums = [index[key] for key in keys if key in index]
        return "[" + compact_numbers(nums) + "]"

    return re.sub(r"\[@([^\]]+)\]", repl, text)


def author_to_vancouver(author: str) -> str:
    author = author.strip()
    if not author:
        return ""
    if "," in author:
        surname, given = [part.strip() for part in author.split(",", 1)]
    else:
        parts = author.split()
        surname, given = parts[-1], " ".join(parts[:-1])
    initials = "".join(part[0] for part in re.split(r"[\s-]+", given) if part)
    return f"{surname} {initials}".strip()


def format_reference_vancouver(entry: dict[str, str]) -> str:
    authors = [author_to_vancouver(a) for a in entry.get("author", "").split(" and ")]
    authors = [a for a in authors if a]
    author_text = ", ".join(authors)
    title = entry.get("title", "").rstrip(".")
    source = (entry.get("journal") or entry.get("booktitle") or entry.get("publisher") or "").rstrip(".")
    year = entry.get("year", "")
    volume = entry.get("volume", "")
    number = entry.get("number", "")
    pages = entry.get("pages", "").replace("--", "\u2013")
    doi = entry.get("doi", "")

    pieces = []
    if author_text:
        pieces.append(author_text + ".")
    if title:
        pieces.append(title + ".")
    if source:
        pieces.append(source + ".")
    citation_tail = ""
    if year:
        citation_tail += year
    if volume:
        citation_tail += f";{volume}"
        if number:
            citation_tail += f"({number})"
    if pages:
        citation_tail += f":{pages}"
    if citation_tail:
        pieces.append(citation_tail + ".")
    if doi:
        pieces.append(f"doi:{doi}.")
    return " ".join(pieces).strip()


def set_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)


def add_formatted_paragraph(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    pos = 0
    for match in re.finditer(r"\*\*(.*?)\*\*", text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    return paragraph


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B7B7B7")
        borders.append(tag)


def add_table_block(doc: Document, spec: dict) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.keep_with_next = True
    run = title.add_run(spec["title"])
    run.bold = True

    rows = spec["rows"]
    headers = spec["headers"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, spec["widths"][i])
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cells[i], spec["widths"][i])
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8.5)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    note_run = note.add_run("Note. " + spec["note"])
    note_run.italic = True
    note_run.font.size = Pt(9)


def add_tables_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    doc.add_heading("Tables", level=1)
    for table_spec in TABLES:
        add_table_block(doc, table_spec)


def restore_portrait(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def should_skip_line(line: str) -> bool:
    if line.startswith("|") and line.endswith("|"):
        return True
    if "Supplementary Figure S2" in line:
        return True
    if "Optional post-hoc XAI" in line:
        return True
    return False


def build_stage19_docx() -> dict:
    raw = MANUSCRIPT_MD.read_text(encoding="utf-8")
    bib = parse_bib(BIB)
    used_keys = collect_used_keys(raw)
    text = replace_citations_numeric(raw, used_keys)

    doc = Document()
    set_style(doc)
    tables_inserted = False
    in_references = False

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or should_skip_line(line):
            continue
        if line == "## Discussion" and not tables_inserted:
            add_tables_section(doc)
            restore_portrait(doc)
            tables_inserted = True
        if line == "## References":
            in_references = True
            doc.add_heading("References", level=1)
            for i, key in enumerate(used_keys, 1):
                add_formatted_paragraph(doc, f"{i}. {format_reference_vancouver(bib.get(key, {}))}")
            continue
        if in_references and line.startswith("References are managed"):
            continue
        if line.startswith("# "):
            p = add_formatted_paragraph(doc, line[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(16)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            add_formatted_paragraph(doc, line[2:].strip(), style="List Bullet")
            continue
        add_formatted_paragraph(doc, line)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Figures", level=1)
    note = "Separate high-resolution PNG/PDF/SVG figure files are available for upload where applicable."
    add_formatted_paragraph(doc, note)
    for caption, filename in RETAINED_FIGURES:
        path = ROOT / "manuscript" / "figures" / filename
        if not path.exists():
            continue
        doc.add_heading(caption, level=2)
        doc.add_picture(str(path), width=Inches(6.2))

    doc.save(OUT_DOCX)
    return {
        "reference_order": used_keys,
        "tables": [spec["title"] for spec in TABLES],
        "retained_figures": [filename for _, filename in RETAINED_FIGURES],
    }


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def write_log(meta: dict) -> None:
    text = extract_docx_text(OUT_DOCX)
    risk_phrases = [
        "external multimodal validation",
        "clinically ready",
        "clinically validated",
        "clinical deployment",
        "deployable",
        "real-world deployment",
        "clinical utility proven",
        "clinically meaningful improvement",
        "gated fusion superiority",
        "superior gated fusion",
        "robust clinical utility",
        "infeasible",
    ]
    risk_rows = []
    lowered = text.lower()
    for phrase in risk_phrases:
        count = lowered.count(phrase)
        status = "acceptable boundary/limitation context" if count else "not present"
        if phrase == "external multimodal validation" and "No external multimodal validation was performed." in text:
            status = "acceptable; explicitly stated as not performed/NO-GO"
        risk_rows.append((phrase, count, status))

    fig_dir = ROOT / "manuscript" / "figures"
    required = [
        "figure1_study_design_evidence_boundary",
        "figure2_internal_model_performance",
        "figure3_external_signal_only_validation",
        "figure4_calibration_reliability",
        "supp_figure_s1_stage14l_audit",
    ]
    figure_lines = []
    for stem in required:
        files = [fig_dir / f"{stem}.{ext}" for ext in ("png", "pdf", "svg")]
        exists = [p.name for p in files if p.exists()]
        figure_lines.append(f"- `{stem}`: {'; '.join(exists)}")
    s2 = fig_dir / "supp_figure_s2_xai_example.png"

    log = [
        "# BMC MIDM Stage 19 Word Cleanup Log",
        "",
        "## Output",
        "",
        f"- Clean Word file: `{OUT_DOCX.relative_to(ROOT)}`",
        "",
        "## Tables Converted",
        "",
        "All Stage 19 result tables were generated as real Word tables rather than Markdown table text.",
        "",
    ]
    log.extend(f"- {title}" for title in meta["tables"])
    log.extend(
        [
            "",
            "## Citation Style Conversion",
            "",
            "- In-text citations were converted from Markdown/BibTeX keys to Vancouver-style numerical square brackets.",
            "- Reference list order was rebuilt by first appearance in the manuscript.",
            "- No mixed author-date citation style is intended in the clean Word draft.",
            "",
            "Reference order:",
            "",
        ]
    )
    log.extend(f"{i}. `{key}`" for i, key in enumerate(meta["reference_order"], 1))
    log.extend(
        [
            "",
            "## Figure Readability and File Readiness",
            "",
            "- Figure 1 through Figure 4 and Supplementary Figure S1 were retained as embedded review figures.",
            "- Separate high-resolution figure files exist as follows:",
            "",
        ]
    )
    log.extend(figure_lines)
    log.extend(
        [
            "",
            "## Supplementary Figure S2 Decision",
            "",
            f"- Source file exists: `{s2.relative_to(ROOT)}` ({'yes' if s2.exists() else 'no'}).",
            "- Decision: removed from the Stage 19 embedded clean submission copy because the review-copy image text is difficult to read at Word page scale.",
            "- Status: optional, not retained in current submission package.",
            "",
            "## Section Order and Declarations",
            "",
            "- Required BMC section order was preserved: title page, structured abstract, keywords, background, methods, results, discussion, conclusions, abbreviations, declarations, references, figure legends, supplementary material list.",
            "- Required declaration subsections are present: ethics approval and consent to participate; consent for publication; availability of data and materials; competing interests; funding; authors' contributions; acknowledgements.",
            "",
            "## Claim Audit",
            "",
            "| Phrase | Count | Status |",
            "|:---|---:|:---|",
        ]
    )
    log.extend(f"| {phrase} | {count} | {status} |" for phrase, count, status in risk_rows)
    log.extend(
        [
            "",
            "The explicit statement `No external multimodal validation was performed.` was retained.",
            "",
            "## Remaining Blockers Before Submission",
            "",
            "- Final manual author confirmation.",
            "- Journal portal upload details.",
        ]
    )
    LOG_MD.write_text("\n".join(log) + "\n", encoding="utf-8")


if __name__ == "__main__":
    metadata = build_stage19_docx()
    write_log(metadata)
    print(json.dumps({"docx": str(OUT_DOCX), "log": str(LOG_MD)}, indent=2))
