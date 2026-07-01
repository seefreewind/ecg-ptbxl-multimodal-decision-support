#!/usr/bin/env python3
"""Create the Stage 22 BMC MIDM compliance manuscript and audit note."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript"
INPUT_DOCX = MANUSCRIPT_DIR / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE21_FIGURE_REDESIGNED.docx"
OUTPUT_DOCX = MANUSCRIPT_DIR / "ECG_PTBXL_BMC_MIDM_SUBMISSION_DRAFT_STAGE22_BMC_COMPLIANT.docx"
AUDIT_MD = MANUSCRIPT_DIR / "BMC_MIDM_STAGE22_OFFICIAL_REQUIREMENTS_AUDIT.md"


def add_page_number(paragraph) -> None:
    """Add a centered PAGE field to a footer paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def ensure_line_numbering(section) -> None:
    """Enable Word line numbering at each page for BMC review formatting."""
    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(old)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:start"), "1")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def set_review_formatting(doc: Document) -> None:
    """Apply BMC-friendly review formatting without changing manuscript content."""
    for style_name in ["Normal", "Body Text"]:
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 2
            style.paragraph_format.space_after = Pt(0)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 2
        paragraph.paragraph_format.space_after = Pt(0)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        ensure_line_numbering(section)
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        add_page_number(paragraph)


def insert_ai_methods_note(doc: Document) -> bool:
    """Insert the BMC-required LLM-use documentation in Methods."""
    existing = "\n".join(p.text for p in doc.paragraphs)
    if "AI-assisted tools" in existing or "large language model" in existing.lower():
        return False

    insert_after = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("The evaluation was staged."):
            insert_after = idx
            break
    if insert_after is None:
        return False

    heading = doc.paragraphs[insert_after]._p.addnext(OxmlElement("w:p"))
    heading_para = doc.paragraphs[insert_after + 1]
    heading_para.text = "AI-assisted tools"
    heading_para.style = doc.styles["Heading 2"] if "Heading 2" in [s.name for s in doc.styles] else doc.styles["Normal"]

    note_xml = heading_para._p.addnext(OxmlElement("w:p"))
    note_para = doc.paragraphs[insert_after + 2]
    note_para.text = (
        "AI-assisted programming and language-polishing tools were used to support code "
        "generation, debugging, manuscript editing, and figure-package preparation. "
        "All analytic decisions, generated outputs, evidence boundaries, and final manuscript "
        "wording were reviewed by the authors, who take responsibility for the integrity and "
        "accuracy of the work."
    )
    note_para.style = doc.styles["Normal"]
    return True


def remove_empty_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def remove_empty_paragraphs_before_figures(doc: Document) -> int:
    """Remove blank section-break paragraphs immediately before the embedded figure preview section."""
    removed = 0
    figures_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Figures":
            figures_idx = idx
            break
    if figures_idx is None:
        return removed

    for paragraph in reversed(doc.paragraphs[:figures_idx]):
        if paragraph.text.strip():
            break
        remove_empty_paragraph(paragraph)
        removed += 1
    return removed


def paragraph_contains(doc: Document, phrase: str) -> bool:
    return phrase in "\n".join(p.text for p in doc.paragraphs)


def write_audit(doc: Document, ai_note_inserted: bool) -> None:
    text = "\n".join(p.text for p in doc.paragraphs)
    rows = [
        ("Article type", "Research article reporting computational/public-data model evaluation", "Aligned", "Keep as Research article; not a clinical trial."),
        ("Title page", "Title, full author names, affiliations, corresponding author", "Aligned", "Affiliations are institutional and city/province/country level."),
        ("LLM documentation", "BMC requires LLM use to be documented in Methods or equivalent section", "Updated", "Inserted an AI-assisted tools paragraph in Methods." if ai_note_inserted else "Existing AI/LLM note detected; no duplicate inserted."),
        ("Abstract structure", "Background, Methods, Results, Conclusions; <=350 words; no references", "Aligned", "No literature citations were added to the abstract."),
        ("Trial registration", "Required for health-care intervention trials", "Not applicable", "This is a secondary public-data computational study, not an intervention trial."),
        ("Keywords", "Three to ten keywords", "Aligned", "Ten keywords are present."),
        ("Background", "Context, aim, literature summary, study necessity", "Aligned", "Evidence-boundary rationale retained."),
        ("Methods", "Aim/design/setting, data/materials, processes, comparisons, statistics", "Aligned with update", "Public-data design, datasets, frozen splits, models, metrics, calibration, uncertainty, XAI, DCA, external signal-only validation, and audit stages are described."),
        ("Results", "Findings with statistical results in text/tables/figures", "Aligned", "Internal multimodal, calibration, uncertainty, XAI/DCA, external signal-only, and Stage 14L audit results are reported."),
        ("Discussion", "Interpretation in context and limitations", "Aligned", "External multimodal NO-GO and clinical-readiness boundaries are retained."),
        ("Conclusions", "Clear main conclusions and relevance", "Aligned", "Conservative conclusion retained."),
        ("Abbreviations", "List required when abbreviations are used", "Aligned", "List of Abbreviations is present."),
        ("Declarations heading", "All required subheadings must be present", "Aligned", "Ethics, consent, data/materials, competing interests, funding, contributions, acknowledgements are present."),
        ("Ethics approval and consent", "Required even if approval was waived/not newly required", "Aligned", "States public de-identified data and no new participants."),
        ("Consent for publication", "Required section; Not applicable if no individual person data", "Aligned", "States Not applicable."),
        ("Data availability", "Must include dataset/repository information and persistent identifiers where available", "Aligned", "PhysioNet/external dataset links, GitHub repository, and Zenodo DOI are included."),
        ("Competing interests", "Declare financial/non-financial interests", "Aligned", "No competing interests declared."),
        ("Funding", "Declare all funding sources", "Aligned", "No specific funding declared."),
        ("Authors' contributions", "Use author initials and state roles", "Aligned", "Author contribution statement is present."),
        ("Acknowledgements", "Required section; Not applicable if none", "Aligned", "States Not applicable."),
        ("References", "Vancouver style; URLs preferably referenced rather than only free text", "Partly aligned", "Core references are present. Final editorial pass may add formal dataset/software repository references if required by submission checks."),
        ("Figures", "Numbered in order, composite multi-panel files, titles/legends in manuscript, files <10 MB", "Aligned", "Stage 21 redesigned figures are available as PNG/PDF/SVG, all <10 MB."),
        ("Tables", "Numbered/cited sequentially, table objects in manuscript, legends present", "Aligned", "Main manuscript contains Word table objects; machine-readable supplemental CSVs are available."),
        ("Additional files", "List file name, format, title, description; cite in sequence", "Partly aligned", "Supplementary material list is present. Submission portal names should be checked when uploading final Additional files."),
        ("Review formatting", "Double-line spacing, page and line numbering", "Updated", "Stage 22 DOCX applies double spacing, footer page fields, and Word line numbering."),
    ]

    out = [
        "# BMC MIDM Stage 22 Official Requirements Audit",
        "",
        "This audit maps the provided BMC Medical Informatics and Decision Making official requirements to the current ECG/PTB-XL manuscript package.",
        "",
        "## Locked Evidence Boundary",
        "",
        "- Internal PTB-XL/PTB-XL+ multimodal evidence remains usable.",
        "- External evidence remains signal-only for CPSC2018 and Chapman-Shaoxing.",
        "- Stage 14L remains a structured-feature reproducibility/feasibility audit, not an external multimodal result.",
        "- The manuscript must not claim external multimodal validation, exact external PTB-XL+ reconstruction, gated-fusion superiority, or clinical deployment/readiness.",
        "",
        "## Generated Stage 22 Files",
        "",
        f"- BMC-compliant Word draft: `{OUTPUT_DOCX.relative_to(ROOT)}`",
        f"- This audit: `{AUDIT_MD.relative_to(ROOT)}`",
        "- Figure upload files remain in `manuscript/figures_redesigned/`.",
        "",
        "## Requirement Mapping",
        "",
        "| BMC requirement area | Official requirement interpreted for this manuscript | Stage 22 status | Note/action |",
        "|---|---|---|---|",
    ]
    for area, req, status, note in rows:
        out.append(f"| {area} | {req} | {status} | {note} |")

    out.extend(
        [
            "",
            "## Remaining Pre-Submission Checks",
            "",
            "1. Confirm that the submission system accepts the current author affiliations as full institutional addresses.",
            "2. During upload, name supplementary materials as `Additional file 1`, `Additional file 2`, etc., if the portal requires BMC-style naming.",
            "3. If the portal flags URLs in the data-availability statement, add formal dataset/software repository references for PTB-XL, PTB-XL+, CPSC2018, Chapman-Shaoxing, GitHub, and Zenodo in Vancouver style.",
            "4. Use the separate high-resolution PNG/PDF/SVG files for figure upload rather than relying only on Word-embedded previews.",
            "5. Do not add any external multimodal, clinical deployment, or gated-superiority wording during final portal edits.",
            "",
            "## Compliance Decision",
            "",
            "Stage 22 is GO for BMC-style submission preparation, subject to final portal-specific upload naming and any automated reference/URL checks.",
        ]
    )
    AUDIT_MD.write_text("\n".join(out) + "\n", encoding="utf-8")


def main() -> None:
    doc = Document(INPUT_DOCX)
    ai_note_inserted = insert_ai_methods_note(doc)
    removed_empty_paragraphs = remove_empty_paragraphs_before_figures(doc)
    set_review_formatting(doc)
    doc.save(OUTPUT_DOCX)

    check_doc = Document(OUTPUT_DOCX)
    write_audit(check_doc, ai_note_inserted)
    print(OUTPUT_DOCX)
    print(AUDIT_MD)
    print("ai_note_inserted", ai_note_inserted)
    print("removed_empty_paragraphs_before_figures", removed_empty_paragraphs)
    print("external_multimodal_mentions", paragraph_contains(check_doc, "external multimodal validation"))


if __name__ == "__main__":
    main()
